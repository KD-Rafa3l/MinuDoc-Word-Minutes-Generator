import tkinter as tk
from tkinter import ttk, filedialog, messagebox, scrolledtext, simpledialog
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import re
from datetime import datetime
import json
from pathlib import Path

class ScrollableFrame(ttk.Frame):
    """Frame scrollable vertical y horizontalmente"""
    def __init__(self, container, *args, **kwargs):
        super().__init__(container, *args, **kwargs)
        
        # Crear canvas y scrollbars
        self.canvas = tk.Canvas(self, highlightthickness=0)
        self.v_scrollbar = ttk.Scrollbar(self, orient="vertical", command=self.canvas.yview)
        self.h_scrollbar = ttk.Scrollbar(self, orient="horizontal", command=self.canvas.xview)
        
        # Configurar canvas
        self.canvas.configure(yscrollcommand=self.v_scrollbar.set, xscrollcommand=self.h_scrollbar.set)
        
        # Frame interior que contendrá todos los widgets
        self.scrollable_frame = ttk.Frame(self.canvas)
        self.scrollable_frame.bind(
            "<Configure>",
            lambda e: self.canvas.configure(scrollregion=self.canvas.bbox("all"))
        )
        
        # Crear ventana en el canvas para el frame scrollable
        self.canvas_frame = self.canvas.create_window((0, 0), window=self.scrollable_frame, anchor="nw")
        
        # Configurar el cambio de tamaño
        self.scrollable_frame.bind("<Configure>", self._on_frame_configure)
        self.canvas.bind("<Configure>", self._on_canvas_configure)
        
        # Empaquetar widgets
        self.canvas.pack(side="left", fill="both", expand=True)
        self.v_scrollbar.pack(side="right", fill="y")
        self.h_scrollbar.pack(side="bottom", fill="x")
        
    def _on_frame_configure(self, event=None):
        """Actualizar scrollregion cuando cambia el tamaño del frame"""
        self.canvas.configure(scrollregion=self.canvas.bbox("all"))
        
    def _on_canvas_configure(self, event):
        """Ajustar el ancho del frame interior al canvas"""
        self.canvas.itemconfig(self.canvas_frame, width=event.width)

class SistemaPlantillasPersonalizadas:
    def __init__(self):
        self.root = tk.Tk()
        self.root.title("Sistema de Plantillas para Minutas Jurídicas - Versión Mejorada")
        self.root.geometry("1400x900")
        self.setup_icon()
        
        # Configurar ventana redimensionable
        self.root.resizable(True, True)
        self.root.minsize(1200, 700)
        
        # Variables de estado
        self.plantillas_personalizadas = {}
        self.plantilla_activa = None
        
        # Crear carpeta de plantillas
        self.carpeta_plantillas = Path("plantillas_personalizadas")
        self.carpeta_plantillas.mkdir(exist_ok=True)
        
        self.configurar_interfaz()
        self.cargar_plantillas_guardadas()
    
    def setup_icon(self):
        try:
            self.root.iconbitmap("law_icon.ico")
        except:
            pass
    
    def configurar_interfaz(self):
        # Configurar estilo
        self.configurar_estilos()
        
        # Frame principal con scroll
        self.main_scrollable = ScrollableFrame(self.root)
        self.main_scrollable.pack(fill="both", expand=True)
        
        # Contenido principal dentro del frame scrollable
        main_content = ttk.Frame(self.main_scrollable.scrollable_frame, padding="20")
        main_content.pack(fill="both", expand=True)
        
        # Header
        header_frame = ttk.Frame(main_content)
        header_frame.pack(fill="x", pady=(0, 20))
        
        titulo = ttk.Label(header_frame, 
                          text="⚖️ SISTEMA DE PLANTILLAS PARA MINUTAS JURÍDICAS", 
                          font=("Arial", 18, "bold"),
                          foreground="#2c3e50")
        titulo.pack(pady=10)
        
        subtitulo = ttk.Label(header_frame,
                             text="Gestión completa de plantillas personalizadas para documentos jurídicos",
                             font=("Arial", 11),
                             foreground="#7f8c8d")
        subtitulo.pack()
        
        # Panel de herramientas rápidas
        quick_tools = ttk.LabelFrame(main_content, text="Acciones Rápidas", padding="15")
        quick_tools.pack(fill="x", pady=(0, 15))
        
        # Grid de botones principales
        tools_grid = ttk.Frame(quick_tools)
        tools_grid.pack(fill="x")
        
        # Fila 1
        ttk.Button(tools_grid, 
                  text="🆕 Crear Nueva Plantilla", 
                  command=self.crear_plantilla_desde_minuta,
                  width=22).grid(row=0, column=0, padx=5, pady=5)
        
        ttk.Button(tools_grid, 
                  text="📝 Editar Plantilla", 
                  command=self.editar_plantilla,
                  width=20).grid(row=0, column=1, padx=5, pady=5)
        
        ttk.Button(tools_grid, 
                  text="🔄 Generar Minuta", 
                  command=self.generar_minuta,
                  width=20).grid(row=0, column=2, padx=5, pady=5)
        
        ttk.Button(tools_grid, 
                  text="📥 Importar", 
                  command=self.importar_plantilla,
                  width=15).grid(row=0, column=3, padx=5, pady=5)
        
        # Fila 2
        ttk.Button(tools_grid, 
                  text="📤 Exportar", 
                  command=self.exportar_plantilla,
                  width=15).grid(row=1, column=0, padx=5, pady=5)
        
        ttk.Button(tools_grid, 
                  text="👁️ Ver Detalles", 
                  command=self.ver_detalles_plantilla,
                  width=15).grid(row=1, column=1, padx=5, pady=5)
        
        ttk.Button(tools_grid, 
                  text="🧹 Limpiar Todo", 
                  command=self.limpiar_formulario,
                  width=15).grid(row=1, column=2, padx=5, pady=5)
        
        ttk.Button(tools_grid, 
                  text="🗑️ Eliminar", 
                  command=self.eliminar_plantilla_activa,
                  width=15).grid(row=1, column=3, padx=5, pady=5)
        
        # Panel de control de plantillas
        control_frame = ttk.LabelFrame(main_content, text="Control de Plantillas Activas", padding="15")
        control_frame.pack(fill="x", pady=(0, 15))
        
        control_grid = ttk.Frame(control_frame)
        control_grid.pack(fill="x")
        
        # Selector de plantilla
        ttk.Label(control_grid, 
                 text="Plantilla Activa:", 
                 font=("Arial", 11, "bold")).grid(row=0, column=0, sticky="w", padx=(0, 10))
        
        self.combo_plantillas = ttk.Combobox(control_grid, 
                                            width=35, 
                                            state="readonly", 
                                            font=("Arial", 10))
        self.combo_plantillas.grid(row=0, column=1, sticky="w", padx=(0, 20))
        self.combo_plantillas.bind('<<ComboboxSelected>>', self.cambiar_plantilla)
        
        # Información de la plantilla
        ttk.Label(control_grid, 
                 text="Descripción:", 
                 font=("Arial", 11)).grid(row=0, column=2, sticky="w", padx=(0, 10))
        
        self.label_descripcion = ttk.Label(control_grid, 
                                          text="Ninguna plantilla seleccionada", 
                                          foreground="#2980b9", 
                                          font=("Arial", 10))
        self.label_descripcion.grid(row=0, column=3, sticky="w")
        
        # Área de trabajo principal con pestañas
        notebook_frame = ttk.Frame(main_content)
        notebook_frame.pack(fill="both", expand=True)
        
        self.notebook = ttk.Notebook(notebook_frame)
        self.notebook.pack(fill="both", expand=True)
        
        # Pestaña 1: Formulario de Datos
        self.tab_formulario = ttk.Frame(self.notebook, padding="10")
        self.notebook.add(self.tab_formulario, text="📋 Formulario de Datos")
        self.configurar_tab_formulario()
        
        # Pestaña 2: Vista Previa
        self.tab_vista_previa = ttk.Frame(self.notebook, padding="10")
        self.notebook.add(self.tab_vista_previa, text="👁️ Vista Previa")
        self.configurar_tab_vista_previa()
        
        # Pestaña 3: Gestión de Plantillas
        self.tab_plantillas = ttk.Frame(self.notebook, padding="10")
        self.notebook.add(self.tab_plantillas, text="📁 Gestión de Plantillas")
        self.configurar_tab_plantillas()
        
        # Barra de estado
        status_frame = ttk.Frame(main_content)
        status_frame.pack(fill="x", pady=(15, 0))
        
        self.status_var = tk.StringVar(value="Sistema listo - Seleccione o cree una plantilla para comenzar")
        status_bar = ttk.Label(status_frame, 
                              textvariable=self.status_var, 
                              relief=tk.SUNKEN, 
                              font=("Arial", 9),
                              background="#f8f9fa")
        status_bar.pack(fill="x")
    
    def configurar_estilos(self):
        """Configurar estilos visuales"""
        style = ttk.Style()
        style.configure("TLabel", font=("Arial", 10))
        style.configure("TButton", font=("Arial", 10))
        style.configure("TEntry", font=("Arial", 10))
        style.configure("TCombobox", font=("Arial", 10))
    
    def configurar_tab_formulario(self):
        # Frame principal con scroll
        form_scrollable = ScrollableFrame(self.tab_formulario)
        form_scrollable.pack(fill="both", expand=True)
        
        form_content = ttk.Frame(form_scrollable.scrollable_frame)
        form_content.pack(fill="both", expand=True, padx=10, pady=10)
        
        # Información de la plantilla
        info_frame = ttk.LabelFrame(form_content, text="Información de la Plantilla Activa", padding="15")
        info_frame.pack(fill="x", pady=(0, 15))
        
        info_grid = ttk.Frame(info_frame)
        info_grid.pack(fill="x")
        
        self.label_info_nombre = ttk.Label(info_grid, text="Nombre: -", font=("Arial", 11, "bold"))
        self.label_info_nombre.grid(row=0, column=0, sticky="w", pady=5, padx=(0, 30))
        
        self.label_info_campos = ttk.Label(info_grid, text="Campos: 0", font=("Arial", 10))
        self.label_info_campos.grid(row=0, column=1, sticky="w", pady=5, padx=(0, 30))
        
        self.label_info_desc = ttk.Label(info_grid, text="Descripción: -", font=("Arial", 10))
        self.label_info_desc.grid(row=0, column=2, sticky="w", pady=5)
        
        # Área de campos del formulario
        campos_frame = ttk.LabelFrame(form_content, text="Campos a Completar", padding="15")
        campos_frame.pack(fill="both", expand=True)
        
        # Frame para campos (sin scroll adicional aquí, ya que el padre tiene scroll)
        self.frame_campos = ttk.Frame(campos_frame)
        self.frame_campos.pack(fill="both", expand=True, padx=10, pady=10)
        
        # Mensaje inicial
        self.label_form_vacio = ttk.Label(self.frame_campos, 
                                         text="Seleccione una plantilla para cargar el formulario correspondiente\n\n"
                                              "Use el botón 'Crear Nueva Plantilla' para comenzar",
                                         font=("Arial", 12), 
                                         foreground="gray", 
                                         justify="center")
        self.label_form_vacio.pack(pady=80)
    
    def configurar_tab_vista_previa(self):
        # Frame principal
        main_preview_frame = ttk.Frame(self.tab_vista_previa)
        main_preview_frame.pack(fill="both", expand=True, padx=10, pady=10)
        
        # Controles de vista previa
        preview_controls = ttk.Frame(main_preview_frame)
        preview_controls.pack(fill="x", pady=(0, 10))
        
        ttk.Label(preview_controls, 
                 text="Vista Previa de la Minuta Generada", 
                 font=("Arial", 14, "bold")).pack(side="left")
        
        ttk.Button(preview_controls, 
                  text="🖨️ Generar Documento Word", 
                  command=self.generar_minuta,
                  width=20).pack(side="right")
        
        # Área de texto para vista previa
        preview_frame = ttk.LabelFrame(main_preview_frame, text="Contenido de la Minuta", padding="10")
        preview_frame.pack(fill="both", expand=True)
        
        self.texto_vista_previa = scrolledtext.ScrolledText(
            preview_frame, 
            wrap=tk.WORD, 
            font=("Courier New", 11),
            padx=10,
            pady=10
        )
        self.texto_vista_previa.pack(fill="both", expand=True)
        self.texto_vista_previa.insert(tk.END, "Complete el formulario y genere la minuta para ver la vista previa aquí...")
    
    def configurar_tab_plantillas(self):
        # Frame principal
        main_management_frame = ttk.Frame(self.tab_plantillas)
        main_management_frame.pack(fill="both", expand=True, padx=10, pady=10)
        
        # Lista de plantillas
        lista_frame = ttk.LabelFrame(main_management_frame, text="Plantillas Disponibles", padding="15")
        lista_frame.pack(fill="both", expand=True, pady=(0, 15))
        
        # Controles de lista
        list_controls = ttk.Frame(lista_frame)
        list_controls.pack(fill="x", pady=(0, 10))
        
        ttk.Label(list_controls, 
                 text="Seleccione una plantilla para gestionar:", 
                 font=("Arial", 11)).pack(side="left")
        
        # Frame para lista y scroll
        list_container = ttk.Frame(lista_frame)
        list_container.pack(fill="both", expand=True)
        
        self.lista_plantillas = tk.Listbox(list_container, height=12, font=("Arial", 11))
        self.lista_plantillas.pack(side="left", fill="both", expand=True)
        
        scroll_lista = ttk.Scrollbar(list_container, orient="vertical", command=self.lista_plantillas.yview)
        scroll_lista.pack(side="right", fill="y")
        self.lista_plantillas.configure(yscrollcommand=scroll_lista.set)
        
        # Botones de gestión
        botones_frame = ttk.Frame(lista_frame)
        botones_frame.pack(fill="x", pady=(10, 0))
        
        ttk.Button(botones_frame, 
                  text="👁️ Ver Detalles Completos", 
                  command=self.ver_detalles_plantilla,
                  width=20).pack(side="left", padx=(0, 10))
        
        ttk.Button(botones_frame, 
                  text="📝 Editar Plantilla", 
                  command=self.editar_plantilla,
                  width=15).pack(side="left", padx=(0, 10))
        
        ttk.Button(botones_frame, 
                  text="📊 Probar Plantilla", 
                  command=self.probar_plantilla,
                  width=15).pack(side="left")
        
        # Panel de detalles
        detalles_frame = ttk.LabelFrame(main_management_frame, text="Detalles de la Plantilla Seleccionada", padding="15")
        detalles_frame.pack(fill="x")
        
        self.texto_detalles = scrolledtext.ScrolledText(
            detalles_frame, 
            wrap=tk.WORD, 
            height=8,
            font=("Arial", 10),
            padx=10,
            pady=10
        )
        self.texto_detalles.pack(fill="both", expand=True)
        self.texto_detalles.insert(tk.END, "Seleccione una plantilla de la lista para ver sus detalles completos...")

    # ===== MÉTODOS DE FUNCIONALIDAD (MANTENIDOS) =====
    
    def crear_plantilla_desde_minuta(self):
        archivo = filedialog.askopenfilename(
            title="Seleccionar minuta base para crear plantilla",
            filetypes=[("Documentos Word", "*.docx"), ("Todos los archivos", "*.*")]
        )
        
        if archivo:
            try:
                doc = Document(archivo)
                contenido = ""
                
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        contenido += paragraph.text + "\n\n"
                
                editor = EditorPlantillasDesdeMinuta(self.root, self.carpeta_plantillas, contenido, archivo)
                self.root.wait_window(editor.ventana)
                self.cargar_plantillas_guardadas()
                
            except Exception as e:
                messagebox.showerror("Error", f"No se pudo cargar la minuta: {str(e)}")
    
    def editar_plantilla(self):
        seleccion = self.lista_plantillas.curselection()
        if seleccion:
            nombre_plantilla = self.lista_plantillas.get(seleccion[0])
            plantilla = self.plantillas_personalizadas.get(nombre_plantilla)
            
            if plantilla:
                editor = EditorPlantillasDesdeMinuta(
                    self.root, self.carpeta_plantillas, 
                    plantilla.get('contenido_base', ''), 
                    plantilla.get('documento_origen', ''),
                    plantilla_existente=plantilla
                )
                self.root.wait_window(editor.ventana)
                self.cargar_plantillas_guardadas()
        else:
            messagebox.showwarning("Advertencia", "Seleccione una plantilla de la lista para editar.")
    
    def generar_minuta(self):
        if not self.plantilla_activa:
            messagebox.showwarning("Advertencia", "No hay plantilla activa. Seleccione una plantilla primero.")
            return
        
        datos = self.obtener_datos_formulario()
        errores = self.validar_formulario(datos)
        
        if errores:
            messagebox.showwarning("Campos requeridos", 
                                "Los siguientes campos son requeridos:\n\n• " + "\n• ".join(errores))
            return
        
        try:
            minuta_generada = self.aplicar_plantilla(self.plantilla_activa, datos)
            
            self.texto_vista_previa.delete("1.0", tk.END)
            self.texto_vista_previa.insert("1.0", minuta_generada)
            
            self.generar_documento_word(minuta_generada)
            
            self.notebook.select(1)
            self.status_var.set("✅ Minuta generada y guardada exitosamente!")
            
        except Exception as e:
            messagebox.showerror("Error", f"No se pudo generar la minuta: {str(e)}")
    
    def obtener_datos_formulario(self):
        datos = {}
        if hasattr(self, 'campos_ui'):
            for campo_id, widget_info in self.campos_ui.items():
                widget = widget_info['widget']
                if isinstance(widget, ttk.Entry):
                    datos[campo_id] = widget.get()
                elif isinstance(widget, tk.Text):
                    datos[campo_id] = widget.get("1.0", tk.END).strip()
                elif isinstance(widget, ttk.Combobox):
                    datos[campo_id] = widget.get()
        return datos
    
    def validar_formulario(self, datos):
        errores = []
        if hasattr(self, 'campos_ui'):
            for campo_id, widget_info in self.campos_ui.items():
                if widget_info.get('requerido', False) and not datos.get(campo_id):
                    errores.append(widget_info['label'])
        return errores
    
    def aplicar_plantilla(self, plantilla, datos):
        contenido_base = plantilla.get('contenido_base', '')
        
        for campo_id, valor in datos.items():
            marcador = f"[[{campo_id}]]"
            contenido_base = contenido_base.replace(marcador, valor)
        
        contenido_base = re.sub(r'\[\[.*?\]\]', '[SIN DATO]', contenido_base)
        return contenido_base
    
    def generar_documento_word(self, contenido):
        doc = Document()
        self.aplicar_formato_apa(doc)
        
        for linea in contenido.split('\n'):
            if linea.strip():
                doc.add_paragraph(linea)
        
        archivo_salida = filedialog.asksaveasfilename(
            title="Guardar minuta como...",
            defaultextension=".docx",
            filetypes=[("Documentos Word", "*.docx")],
            initialfile=f"minuta_{datetime.now().strftime('%Y%m%d_%H%M')}.docx"
        )
        
        if archivo_salida:
            doc.save(archivo_salida)
            os.startfile(archivo_salida)
            return True
        return False
    
    def aplicar_formato_apa(self, doc):
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)
        
        paragraph_format = style.paragraph_format
        paragraph_format.line_spacing = 2.0
    
    def cargar_plantillas_guardadas(self):
        self.plantillas_personalizadas = {}
        for archivo in self.carpeta_plantillas.glob("*.json"):
            try:
                with open(archivo, 'r', encoding='utf-8') as f:
                    plantilla = json.load(f)
                self.plantillas_personalizadas[archivo.stem] = plantilla
            except Exception as e:
                print(f"Error cargando plantilla {archivo}: {e}")
        
        self.actualizar_listas_plantillas()
    
    def actualizar_listas_plantillas(self):
        plantillas = list(self.plantillas_personalizadas.keys())
        self.combo_plantillas['values'] = plantillas
        self.lista_plantillas.delete(0, tk.END)
        
        for nombre in plantillas:
            self.lista_plantillas.insert(tk.END, nombre)
        
        if plantillas:
            self.combo_plantillas.set(plantillas[0])
            self.cambiar_plantilla()
    
    def cambiar_plantilla(self, event=None):
        nombre_plantilla = self.combo_plantillas.get()
        if nombre_plantilla in self.plantillas_personalizadas:
            self.plantilla_activa = self.plantillas_personalizadas[nombre_plantilla]
            self.cargar_formulario_plantilla()
            self.actualizar_info_plantilla()
            self.status_var.set(f"✅ Plantilla activa: {nombre_plantilla}")
    
    def cargar_formulario_plantilla(self):
        for widget in self.frame_campos.winfo_children():
            widget.destroy()
        
        self.campos_ui = {}
        
        if not self.plantilla_activa:
            return
        
        campos = self.plantilla_activa.get('campos_personalizados', [])
        
        if not campos:
            self.label_form_vacio = ttk.Label(self.frame_campos, 
                                             text="Esta plantilla no tiene campos personalizados definidos",
                                             font=("Arial", 11), foreground="gray")
            self.label_form_vacio.pack(pady=50)
            return
        
        for i, campo in enumerate(campos):
            self.crear_campo_formulario(campo, i)
    
    def crear_campo_formulario(self, campo, index):
        frame_campo = ttk.Frame(self.frame_campos)
        frame_campo.pack(fill="x", pady=8, padx=15)
        
        label_text = campo['nombre']
        if campo.get('requerido', False):
            label_text += " *"
        
        label = ttk.Label(frame_campo, text=label_text, width=25, anchor="w", font=("Arial", 10))
        label.pack(side="left", padx=(0, 15))
        
        campo_id = campo['id']
        if campo['tipo'] == 'texto':
            widget = ttk.Entry(frame_campo, width=50, font=("Arial", 9))
            widget.pack(side="left", fill="x", expand=True)
            
        elif campo['tipo'] == 'textarea':
            frame_text = ttk.Frame(frame_campo)
            frame_text.pack(side="left", fill="x", expand=True)
            
            widget = tk.Text(frame_text, width=60, height=4, wrap=tk.WORD, font=("Arial", 9))
            scrollbar = ttk.Scrollbar(frame_text, orient="vertical", command=widget.yview)
            widget.configure(yscrollcommand=scrollbar.set)
            
            widget.pack(side="left", fill="both", expand=True)
            scrollbar.pack(side="right", fill="y")
            
        elif campo['tipo'] == 'seleccion':
            widget = ttk.Combobox(frame_campo, width=48, values=campo.get('opciones', []), font=("Arial", 9))
            widget.pack(side="left", fill="x", expand=True)
            
        elif campo['tipo'] == 'fecha':
            widget = ttk.Entry(frame_campo, width=25, font=("Arial", 9))
            widget.pack(side="left")
            ttk.Label(frame_campo, text="(DD/MM/AAAA)", font=("Arial", 8), foreground="gray").pack(side="left", padx=(5, 0))
        
        if campo.get('descripcion'):
            self.crear_tooltip(label, campo['descripcion'])
        
        self.campos_ui[campo_id] = {
            'widget': widget,
            'label': campo['nombre'],
            'requerido': campo.get('requerido', False)
        }
    
    def crear_tooltip(self, widget, text):
        def on_enter(event):
            tooltip = tk.Toplevel()
            tooltip.wm_overrideredirect(True)
            tooltip.wm_geometry(f"+{event.x_root+10}+{event.y_root+10}")
            label = ttk.Label(tooltip, text=text, background="lightyellow", 
                            relief="solid", borderwidth=1, padding=5, font=("Arial", 9))
            label.pack()
            widget.tooltip = tooltip
        
        def on_leave(event):
            if hasattr(widget, 'tooltip'):
                widget.tooltip.destroy()
        
        widget.bind("<Enter>", on_enter)
        widget.bind("<Leave>", on_leave)
    
    def actualizar_info_plantilla(self):
        if self.plantilla_activa:
            self.label_info_nombre.config(text=f"Nombre: {self.plantilla_activa.get('nombre', 'N/A')}")
            self.label_info_desc.config(text=f"Descripción: {self.plantilla_activa.get('descripcion', 'N/A')}")
            
            campos = self.plantilla_activa.get('campos_personalizados', [])
            campos_requeridos = sum(1 for c in campos if c.get('requerido', False))
            self.label_info_campos.config(text=f"Campos: {len(campos)} (Requeridos: {campos_requeridos})")
            
            self.label_descripcion.config(text=self.plantilla_activa.get('descripcion', 'Sin descripción'))
    
    def importar_plantilla(self):
        archivo = filedialog.askopenfilename(
            title="Importar plantilla",
            filetypes=[("Archivos de plantilla", "*.json")]
        )
        
        if archivo:
            try:
                with open(archivo, 'r', encoding='utf-8') as f:
                    plantilla = json.load(f)
                
                nombre = plantilla.get('nombre', Path(archivo).stem)
                
                if nombre in self.plantillas_personalizadas:
                    respuesta = messagebox.askyesno("Confirmar", 
                                                  f"¿Sobrescribir la plantilla existente '{nombre}'?")
                    if not respuesta:
                        return
                
                archivo_destino = self.carpeta_plantillas / f"{nombre}.json"
                with open(archivo_destino, 'w', encoding='utf-8') as f:
                    json.dump(plantilla, f, ensure_ascii=False, indent=2)
                
                self.cargar_plantillas_guardadas()
                messagebox.showinfo("Éxito", f"Plantilla '{nombre}' importada correctamente.")
                
            except Exception as e:
                messagebox.showerror("Error", f"No se pudo importar la plantilla: {str(e)}")
    
    def exportar_plantilla(self):
        seleccion = self.lista_plantillas.curselection()
        if seleccion:
            nombre_plantilla = self.lista_plantillas.get(seleccion[0])
            plantilla = self.plantillas_personalizadas.get(nombre_plantilla)
            
            if plantilla:
                archivo = filedialog.asksaveasfilename(
                    title="Exportar plantilla como...",
                    defaultextension=".json",
                    filetypes=[("Archivos de plantilla", "*.json")],
                    initialfile=f"{nombre_plantilla}.json"
                )
                
                if archivo:
                    try:
                        with open(archivo, 'w', encoding='utf-8') as f:
                            json.dump(plantilla, f, ensure_ascii=False, indent=2)
                        messagebox.showinfo("Éxito", f"Plantilla exportada a: {archivo}")
                    except Exception as e:
                        messagebox.showerror("Error", f"No se pudo exportar: {str(e)}")
        else:
            messagebox.showwarning("Advertencia", "Seleccione una plantilla para exportar.")
    
    def eliminar_plantilla_activa(self):
        nombre_plantilla = self.combo_plantillas.get()
        if nombre_plantilla and nombre_plantilla in self.plantillas_personalizadas:
            respuesta = messagebox.askyesno("Confirmar", 
                                          f"¿Está seguro de eliminar la plantilla '{nombre_plantilla}'?")
            if respuesta:
                archivo_plantilla = self.carpeta_plantillas / f"{nombre_plantilla}.json"
                if archivo_plantilla.exists():
                    archivo_plantilla.unlink()
                
                self.cargar_plantillas_guardadas()
                messagebox.showinfo("Éxito", f"Plantilla '{nombre_plantilla}' eliminada.")
        else:
            messagebox.showwarning("Advertencia", "No hay plantilla seleccionada para eliminar.")
    
    def probar_plantilla(self):
        seleccion = self.lista_plantillas.curselection()
        if seleccion:
            nombre_plantilla = self.lista_plantillas.get(seleccion[0])
            self.combo_plantillas.set(nombre_plantilla)
            self.cambiar_plantilla()
            self.notebook.select(0)
            messagebox.showinfo("Éxito", f"Plantilla '{nombre_plantilla}' activada para prueba.")
        else:
            messagebox.showwarning("Advertencia", "Seleccione una plantilla de la lista.")
    
    def ver_detalles_plantilla(self):
        seleccion = self.lista_plantillas.curselection()
        if seleccion:
            nombre_plantilla = self.lista_plantillas.get(seleccion[0])
            plantilla = self.plantillas_personalizadas.get(nombre_plantilla)
            
            if plantilla:
                detalles = f"""INFORMACIÓN DETALLADA DE LA PLANTILLA

Nombre: {plantilla.get('nombre', 'N/A')}
Descripción: {plantilla.get('descripcion', 'N/A')}
Tipo: {plantilla.get('tipo', 'N/A')}
Fecha creación: {plantilla.get('fecha_creacion', 'N/A')}
Documento origen: {plantilla.get('documento_origen', 'N/A')}

CAMPOS PERSONALIZADOS:
"""
                campos = plantilla.get('campos_personalizados', [])
                for i, campo in enumerate(campos, 1):
                    requerido = "SÍ" if campo.get('requerido') else "no"
                    detalles += f"\n{i}. {campo['nombre']} ({campo['tipo']}) - Requerido: {requerido}"
                    if campo.get('descripcion'):
                        detalles += f"\n   Descripción: {campo['descripcion']}"
                
                self.texto_detalles.delete("1.0", tk.END)
                self.texto_detalles.insert("1.0", detalles)
        else:
            messagebox.showwarning("Advertencia", "Seleccione una plantilla de la lista.")
    
    def limpiar_formulario(self):
        if hasattr(self, 'campos_ui'):
            for campo_id, widget_info in self.campos_ui.items():
                widget = widget_info['widget']
                if isinstance(widget, ttk.Entry) or isinstance(widget, ttk.Combobox):
                    widget.delete(0, tk.END)
                elif isinstance(widget, tk.Text):
                    widget.delete("1.0", tk.END)
        
        self.texto_vista_previa.delete("1.0", tk.END)
        self.texto_vista_previa.insert(tk.END, "Formulario limpiado. Complete los campos y genere una nueva minuta.")
        self.status_var.set("Formulario limpiado - Listo para nuevo proceso")


class EditorPlantillasDesdeMinuta:
    def __init__(self, parent, carpeta_plantillas, contenido_minuta="", archivo_origen="", plantilla_existente=None):
        self.parent = parent
        self.carpeta_plantillas = carpeta_plantillas
        self.contenido_minuta = contenido_minuta
        self.archivo_origen = archivo_origen
        self.plantilla_existente = plantilla_existente
        
        self.ventana = tk.Toplevel(parent)
        self.ventana.title("Editor de Plantillas - Crear/Editar Plantilla")
        self.ventana.geometry("1400x900")
        self.ventana.transient(parent)
        self.ventana.grab_set()
        self.ventana.resizable(True, True)
        self.ventana.minsize(1200, 700)
        
        self.campos_personalizados = []
        self.mapeo_selecciones = {}
        self.texto_seleccionado_actual = None
        self.posicion_seleccion_actual = None
        
        # Frame principal con scroll
        self.main_scrollable = ScrollableFrame(self.ventana)
        self.main_scrollable.pack(fill="both", expand=True)
        
        self.configurar_interfaz()
        
        if plantilla_existente:
            self.cargar_plantilla_existente(plantilla_existente)
    
    def configurar_interfaz(self):
        # Contenido principal dentro del frame scrollable
        main_content = ttk.Frame(self.main_scrollable.scrollable_frame, padding="20")
        main_content.pack(fill="both", expand=True)
        
        # Header
        header_frame = ttk.Frame(main_content)
        header_frame.pack(fill="x", pady=(0, 20))
        
        titulo = ttk.Label(header_frame, 
                          text="✏️ Editor de Plantillas - Seleccione Texto para Crear Campos", 
                          font=("Arial", 16, "bold"),
                          foreground="#2c3e50")
        titulo.pack()
        
        subtitulo = ttk.Label(header_frame,
                             text="Seleccione texto en el documento y cree campos personalizados reemplazando con marcadores",
                             font=("Arial", 11),
                             foreground="#7f8c8d")
        subtitulo.pack()
        
        # Información básica
        info_frame = ttk.LabelFrame(main_content, text="Información Básica de la Plantilla", padding="15")
        info_frame.pack(fill="x", pady=(0, 15))
        
        # Grid para información
        info_grid = ttk.Frame(info_frame)
        info_grid.pack(fill="x")
        
        ttk.Label(info_grid, text="Nombre de la plantilla:", font=("Arial", 10)).grid(row=0, column=0, sticky="w", pady=8)
        self.entry_nombre = ttk.Entry(info_grid, width=50, font=("Arial", 10))
        self.entry_nombre.grid(row=0, column=1, sticky="w", pady=8, padx=(10, 0))
        
        ttk.Label(info_grid, text="Descripción:", font=("Arial", 10)).grid(row=1, column=0, sticky="w", pady=8)
        self.entry_descripcion = ttk.Entry(info_grid, width=50, font=("Arial", 10))
        self.entry_descripcion.grid(row=1, column=1, sticky="w", pady=8, padx=(10, 0))
        
        ttk.Label(info_grid, text="Tipo:", font=("Arial", 10)).grid(row=2, column=0, sticky="w", pady=8)
        self.combo_tipo = ttk.Combobox(info_grid, width=47, font=("Arial", 10),
                                      values=["Amparo", "Contrato", "Demanda", "Recurso", "General", "Solicitud"])
        self.combo_tipo.grid(row=2, column=1, sticky="w", pady=8, padx=(10, 0))
        self.combo_tipo.set("General")
        
        # Área de trabajo dividida
        workspace_frame = ttk.Frame(main_content)
        workspace_frame.pack(fill="both", expand=True, pady=(0, 15))
        
        # Panel izquierdo - Contenido de la minuta
        left_panel = ttk.LabelFrame(workspace_frame, text="📄 Contenido de la Minuta Base", padding="15")
        left_panel.pack(side="left", fill="both", expand=True, padx=(0, 10))
        
        instrucciones = ttk.Label(left_panel, 
                 text="Seleccione texto y haga clic en 'Crear Campo' para reemplazar con marcadores", 
                 font=("Arial", 9, "italic"))
        instrucciones.pack(anchor="w", pady=(0, 10))
        
        self.texto_minuta = scrolledtext.ScrolledText(
            left_panel, 
            wrap=tk.WORD, 
            font=("Consolas", 10)
        )
        self.texto_minuta.pack(fill="both", expand=True)
        
        if self.contenido_minuta:
            self.texto_minuta.insert("1.0", self.contenido_minuta)
        
        self.texto_minuta.tag_configure("seleccionado", background="lightgreen", foreground="darkgreen")
        self.texto_minuta.bind("<<Selection>>", self.guardar_seleccion_actual)
        
        # Panel derecho - Configuración de campos
        right_panel = ttk.LabelFrame(workspace_frame, text="⚙️ Configuración de Campos Personalizados", padding="15")
        right_panel.pack(side="right", fill="both", expand=True)
        
        # Instrucciones
        instrucciones_frame = ttk.Frame(right_panel)
        instrucciones_frame.pack(fill="x", pady=(0, 15))
        
        ttk.Label(instrucciones_frame, 
                 text="Cómo crear campos personalizados:", 
                 font=("Arial", 11, "bold")).pack(anchor="w")
        
        instrucciones_texto = ttk.Label(instrucciones_frame, 
                                text="1. Seleccione texto en la minuta\n"
                                     "2. Haga clic en 'Crear Campo desde Selección'\n"
                                     "3. Configure las propiedades del campo\n"
                                     "4. El texto se reemplazará automáticamente con [[ID_CAMPO]]",
                                font=("Arial", 9), 
                                justify="left")
        instrucciones_texto.pack(anchor="w", pady=(5, 0))
        
        # Botones de acción
        action_buttons = ttk.Frame(right_panel)
        action_buttons.pack(fill="x", pady=10)
        
        ttk.Button(action_buttons, 
                  text="📍 Crear Campo desde Selección Actual", 
                  command=self.crear_campo_desde_seleccion,
                  width=28).pack(side="left", padx=(0, 10))
        
        ttk.Button(action_buttons, 
                  text="➕ Agregar Campo Manualmente", 
                  command=self.agregar_campo_manual,
                  width=22).pack(side="left")
        
        # Lista de campos creados
        campos_frame = ttk.LabelFrame(right_panel, text="Campos Creados", padding="10")
        campos_frame.pack(fill="both", expand=True, pady=(0, 10))
        
        list_container = ttk.Frame(campos_frame)
        list_container.pack(fill="both", expand=True)
        
        self.lista_campos = tk.Listbox(list_container, height=8, font=("Arial", 10))
        self.lista_campos.pack(side="left", fill="both", expand=True)
        
        scroll_campos = ttk.Scrollbar(list_container, orient="vertical", command=self.lista_campos.yview)
        scroll_campos.pack(side="right", fill="y")
        self.lista_campos.configure(yscrollcommand=scroll_campos.set)
        
        # Botones de gestión de campos
        manage_buttons = ttk.Frame(right_panel)
        manage_buttons.pack(fill="x", pady=5)
        
        ttk.Button(manage_buttons, 
                  text="📝 Editar Campo Seleccionado", 
                  command=self.editar_campo,
                  width=20).pack(side="left", padx=(0, 10))
        
        ttk.Button(manage_buttons, 
                  text="🗑️ Eliminar Campo Seleccionado", 
                  command=self.eliminar_campo,
                  width=20).pack(side="left")
        
        # Botones finales
        final_buttons = ttk.Frame(main_content)
        final_buttons.pack(fill="x", pady=10)
        
        ttk.Button(final_buttons, 
                  text="💾 Guardar Plantilla", 
                  command=self.guardar_plantilla,
                  width=20).pack(side="left", padx=(0, 15))
        
        ttk.Button(final_buttons, 
                  text="🔍 Vista Previa de Marcadores", 
                  command=self.mostrar_vista_previa,
                  width=22).pack(side="left", padx=(0, 15))
        
        ttk.Button(final_buttons, 
                  text="❌ Cancelar y Salir", 
                  command=self.ventana.destroy,
                  width=16).pack(side="left")

    # Los métodos de funcionalidad se mantienen igual...
    def guardar_seleccion_actual(self, event=None):
        try:
            if self.texto_minuta.tag_ranges(tk.SEL):
                self.texto_seleccionado_actual = self.texto_minuta.get(tk.SEL_FIRST, tk.SEL_LAST)
                self.posicion_seleccion_actual = (tk.SEL_FIRST, tk.SEL_LAST)
        except:
            self.texto_seleccionado_actual = None
            self.posicion_seleccion_actual = None
    
    def crear_campo_desde_seleccion(self):
        if not self.texto_seleccionado_actual:
            messagebox.showwarning("Advertencia", "Primero seleccione texto en la minuta.")
            return
        
        texto_seleccionado = self.texto_seleccionado_actual
        dialogo = DialogoCampoDesdeSeleccion(self.ventana, texto_seleccionado)
        self.ventana.wait_window(dialogo.ventana)
        
        if dialogo.campo_creado:
            campo = dialogo.campo_creado
            self.campos_personalizados.append(campo)
            
            if self.posicion_seleccion_actual:
                inicio, fin = self.posicion_seleccion_actual
                marcador = f"[[{campo['id']}]]"
                
                self.texto_minuta.delete(inicio, fin)
                self.texto_minuta.insert(inicio, marcador)
                
                nuevo_fin = self.texto_minuta.index(f"{inicio} + {len(marcador)}c")
                self.texto_minuta.tag_add("seleccionado", inicio, nuevo_fin)
            
            self.mapeo_selecciones[campo['id']] = {
                'texto_original': texto_seleccionado,
                'marcador': marcador
            }
            
            self.actualizar_lista_campos()
            messagebox.showinfo("Éxito", f"Campo '{campo['nombre']}' creado correctamente.")
            
            self.texto_seleccionado_actual = None
            self.posicion_seleccion_actual = None
    
    def agregar_campo_manual(self):
        dialogo = DialogoCampoDesdeSeleccion(self.ventana, "")
        self.ventana.wait_window(dialogo.ventana)
        
        if dialogo.campo_creado:
            self.campos_personalizados.append(dialogo.campo_creado)
            self.actualizar_lista_campos()
            messagebox.showinfo("Éxito", f"Campo '{dialogo.campo_creado['nombre']}' agregado manualmente.")
    
    def editar_campo(self):
        seleccion = self.lista_campos.curselection()
        if not seleccion:
            messagebox.showwarning("Advertencia", "Seleccione un campo para editar.")
            return
        
        index = seleccion[0]
        campo_existente = self.campos_personalizados[index]
        
        dialogo = DialogoCampoDesdeSeleccion(self.ventana, "", campo_existente)
        self.ventana.wait_window(dialogo.ventana)
        
        if dialogo.campo_creado:
            self.campos_personalizados[index] = dialogo.campo_creado
            self.actualizar_lista_campos()
            messagebox.showinfo("Éxito", f"Campo '{dialogo.campo_creado['nombre']}' actualizado.")
    
    def eliminar_campo(self):
        seleccion = self.lista_campos.curselection()
        if not seleccion:
            messagebox.showwarning("Advertencia", "Seleccione un campo para eliminar.")
            return
        
        index = seleccion[0]
        campo = self.campos_personalizados[index]
        
        respuesta = messagebox.askyesno("Confirmar", 
                                      f"¿Está seguro de eliminar el campo '{campo['nombre']}'?")
        if respuesta:
            if campo['id'] in self.mapeo_selecciones:
                contenido_actual = self.texto_minuta.get("1.0", tk.END)
                marcador = f"[[{campo['id']}]]"
                texto_original = self.mapeo_selecciones[campo['id']]['texto_original']
                nuevo_contenido = contenido_actual.replace(marcador, texto_original)
                
                self.texto_minuta.delete("1.0", tk.END)
                self.texto_minuta.insert("1.0", nuevo_contenido)
                
                del self.mapeo_selecciones[campo['id']]
            
            self.campos_personalizados.pop(index)
            self.actualizar_lista_campos()
            messagebox.showinfo("Éxito", f"Campo '{campo['nombre']}' eliminado.")
    
    def actualizar_lista_campos(self):
        self.lista_campos.delete(0, tk.END)
        for campo in self.campos_personalizados:
            requerido = " *" if campo.get('requerido') else ""
            self.lista_campos.insert(tk.END, f"{campo['nombre']}{requerido} ({campo['tipo']})")
    
    def cargar_plantilla_existente(self, plantilla):
        self.entry_nombre.delete(0, tk.END)
        self.entry_nombre.insert(0, plantilla.get('nombre', ''))
        
        self.entry_descripcion.delete(0, tk.END)
        self.entry_descripcion.insert(0, plantilla.get('descripcion', ''))
        
        self.combo_tipo.set(plantilla.get('tipo', 'General'))
        
        self.texto_minuta.delete("1.0", tk.END)
        self.texto_minuta.insert("1.0", plantilla.get('contenido_base', ''))
        
        self.campos_personalizados = plantilla.get('campos_personalizados', [])
        self.actualizar_lista_campos()
        
        self.resaltar_marcadores()
    
    def resaltar_marcadores(self):
        contenido = self.texto_minuta.get("1.0", tk.END)
        for marcador in re.findall(r'\[\[.*?\]\]', contenido):
            inicio = "1.0"
            while True:
                inicio = self.texto_minuta.search(marcador, inicio, tk.END)
                if not inicio:
                    break
                fin = f"{inicio} + {len(marcador)}c"
                self.texto_minuta.tag_add("seleccionado", inicio, fin)
                inicio = fin
    
    def mostrar_vista_previa(self):
        if not self.campos_personalizados:
            messagebox.showinfo("Marcadores", "No hay campos creados todavía.")
            return
        
        marcadores = "\n".join([f"[[{campo['id']}]] - {campo['nombre']} ({campo['tipo']})" 
                              for campo in self.campos_personalizados])
        messagebox.showinfo("Marcadores Disponibles", 
                          f"Puede usar estos marcadores en el contenido:\n\n{marcadores}")
    
    def guardar_plantilla(self):
        nombre = self.entry_nombre.get().strip()
        descripcion = self.entry_descripcion.get().strip()
        tipo = self.combo_tipo.get()
        contenido = self.texto_minuta.get("1.0", tk.END).strip()
        
        if not nombre:
            messagebox.showwarning("Advertencia", "El nombre de la plantilla es requerido.")
            return
        
        if not self.campos_personalizados:
            messagebox.showwarning("Advertencia", "Debe crear al menos un campo para la plantilla.")
            return
        
        plantilla = {
            'nombre': nombre,
            'descripcion': descripcion,
            'tipo': tipo,
            'fecha_creacion': datetime.now().isoformat(),
            'campos_personalizados': self.campos_personalizados,
            'contenido_base': contenido,
            'documento_origen': self.archivo_origen
        }
        
        archivo_plantilla = self.carpeta_plantillas / f"{nombre}.json"
        
        try:
            with open(archivo_plantilla, 'w', encoding='utf-8') as f:
                json.dump(plantilla, f, ensure_ascii=False, indent=2)
            
            messagebox.showinfo("Éxito", f"Plantilla '{nombre}' guardada correctamente!")
            self.ventana.destroy()
            
        except Exception as e:
            messagebox.showerror("Error", f"No se pudo guardar la plantilla: {str(e)}")


class DialogoCampoDesdeSeleccion:
    def __init__(self, parent, texto_seleccionado="", campo_existente=None):
        self.parent = parent
        self.texto_seleccionado = texto_seleccionado
        self.campo_existente = campo_existente
        self.campo_creado = None
        
        self.ventana = tk.Toplevel(parent)
        self.ventana.title("Configurar Campo Personalizado")
        self.ventana.geometry("600x650")
        self.ventana.transient(parent)
        self.ventana.grab_set()
        self.ventana.resizable(True, True)
        self.ventana.minsize(550, 600)
        
        # Frame con scroll para el diálogo
        self.dialog_scrollable = ScrollableFrame(self.ventana)
        self.dialog_scrollable.pack(fill="both", expand=True)
        
        self.configurar_interfaz()
        
        if campo_existente:
            self.cargar_datos_existentes(campo_existente)
    
    def configurar_interfaz(self):
        # Contenido principal dentro del frame scrollable
        main_content = ttk.Frame(self.dialog_scrollable.scrollable_frame, padding="20")
        main_content.pack(fill="both", expand=True)
        
        ttk.Label(main_content, 
                 text="⚙️ Configurar Campo Personalizado", 
                 font=("Arial", 14, "bold")).pack(pady=(0, 20))
        
        if self.texto_seleccionado:
            frame_texto = ttk.LabelFrame(main_content, text="Texto Seleccionado", padding="10")
            frame_texto.pack(fill="x", pady=10)
            
            label_texto = ttk.Label(frame_texto, text=self.texto_seleccionado, 
                                   wraplength=500, font=("Arial", 10), background="lightyellow")
            label_texto.pack(fill="x")
        
        # Configuración del campo
        config_frame = ttk.LabelFrame(main_content, text="Propiedades del Campo", padding="15")
        config_frame.pack(fill="both", expand=True, pady=10)
        
        # ID del campo
        ttk.Label(config_frame, text="ID del campo (interno):", font=("Arial", 10)).grid(row=0, column=0, sticky="w", pady=8)
        self.entry_id = ttk.Entry(config_frame, width=30, font=("Arial", 10))
        self.entry_id.grid(row=0, column=1, sticky="w", pady=8, padx=(10, 0))
        ttk.Label(config_frame, text="(sin espacios, único)", font=("Arial", 9), foreground="gray").grid(row=0, column=2, sticky="w", pady=8, padx=(5, 0))
        
        # Nombre del campo
        ttk.Label(config_frame, text="Nombre visible:", font=("Arial", 10)).grid(row=1, column=0, sticky="w", pady=8)
        self.entry_nombre = ttk.Entry(config_frame, width=30, font=("Arial", 10))
        self.entry_nombre.grid(row=1, column=1, sticky="w", pady=8, padx=(10, 0))
        
        if self.texto_seleccionado and not self.campo_existente:
            nombre_sugerido = self.texto_seleccionado.strip()[:30]
            self.entry_nombre.insert(0, nombre_sugerido)
            
            id_sugerido = re.sub(r'[^a-zA-Z0-9_]', '_', nombre_sugerido.lower())
            self.entry_id.insert(0, id_sugerido)
        
        # Tipo de campo
        ttk.Label(config_frame, text="Tipo de campo:", font=("Arial", 10)).grid(row=2, column=0, sticky="w", pady=8)
        
        tipo_frame = ttk.Frame(config_frame)
        tipo_frame.grid(row=2, column=1, columnspan=2, sticky="w", pady=8, padx=(10, 0))
        
        self.tipo_var = tk.StringVar(value="texto")
        
        tk.Radiobutton(tipo_frame, text="Texto corto", variable=self.tipo_var, value="texto", 
                      font=("Arial", 10)).pack(anchor="w", pady=2)
        tk.Radiobutton(tipo_frame, text="Texto largo", variable=self.tipo_var, value="textarea", 
                      font=("Arial", 10)).pack(anchor="w", pady=2)
        tk.Radiobutton(tipo_frame, text="Selección", variable=self.tipo_var, value="seleccion", 
                      font=("Arial", 10)).pack(anchor="w", pady=2)
        tk.Radiobutton(tipo_frame, text="Fecha", variable=self.tipo_var, value="fecha", 
                      font=("Arial", 10)).pack(anchor="w", pady=2)
        
        # Opciones para selección
        self.frame_opciones = ttk.LabelFrame(config_frame, text="Opciones de Selección", padding="10")
        self.frame_opciones.grid(row=3, column=0, columnspan=3, sticky="we", pady=10)
        
        ttk.Label(self.frame_opciones, text="Una opción por línea:", font=("Arial", 9)).pack(anchor="w")
        self.texto_opciones = tk.Text(self.frame_opciones, height=6, width=50, font=("Arial", 9))
        self.texto_opciones.pack(fill="x", pady=5)
        
        # Descripción
        ttk.Label(config_frame, text="Descripción/tooltip:", font=("Arial", 10)).grid(row=4, column=0, sticky="w", pady=8)
        self.entry_descripcion = ttk.Entry(config_frame, width=30, font=("Arial", 10))
        self.entry_descripcion.grid(row=4, column=1, columnspan=2, sticky="we", pady=8, padx=(10, 0))
        
        # Campo requerido
        requerido_frame = ttk.Frame(config_frame)
        requerido_frame.grid(row=5, column=0, columnspan=3, sticky="w", pady=15)
        
        self.requerido_var = tk.BooleanVar(value=True)
        tk.Checkbutton(requerido_frame, text="Campo requerido", 
                      variable=self.requerido_var, font=("Arial", 10)).pack(anchor="w")
        
        # Botones
        botones_frame = ttk.Frame(main_content)
        botones_frame.pack(fill="x", pady=10)
        
        ttk.Button(botones_frame, 
                  text="💾 Guardar Campo", 
                  command=self.guardar_campo,
                  width=16).pack(side="left", padx=(0, 10))
        
        ttk.Button(botones_frame, 
                  text="❌ Cancelar", 
                  command=self.ventana.destroy,
                  width=12).pack(side="left")
    
    def cargar_datos_existentes(self, campo):
        self.entry_id.insert(0, campo.get('id', ''))
        self.entry_id.config(state='disabled')
        
        self.entry_nombre.insert(0, campo.get('nombre', ''))
        self.tipo_var.set(campo.get('tipo', 'texto'))
        self.entry_descripcion.insert(0, campo.get('descripcion', ''))
        self.requerido_var.set(campo.get('requerido', False))
        
        if campo.get('tipo') == 'seleccion' and 'opciones' in campo:
            self.texto_opciones.insert("1.0", "\n".join(campo['opciones']))
    
    def guardar_campo(self):
        campo_id = self.entry_id.get().strip()
        nombre = self.entry_nombre.get().strip()
        tipo = self.tipo_var.get()
        descripcion = self.entry_descripcion.get().strip()
        requerido = self.requerido_var.get()
        
        if not campo_id:
            messagebox.showwarning("Advertencia", "El ID del campo es requerido.")
            return
        
        if not nombre:
            messagebox.showwarning("Advertencia", "El nombre del campo es requerido.")
            return
        
        if ' ' in campo_id:
            messagebox.showwarning("Advertencia", "El ID no puede contener espacios.")
            return
        
        campo = {
            'id': campo_id,
            'nombre': nombre,
            'tipo': tipo,
            'descripcion': descripcion,
            'requerido': requerido
        }
        
        if tipo == 'seleccion':
            opciones_texto = self.texto_opciones.get("1.0", tk.END).strip()
            if opciones_texto:
                campo['opciones'] = [opcion.strip() for opcion in opciones_texto.split('\n') if opcion.strip()]
            else:
                messagebox.showwarning("Advertencia", "Debe proporcionar opciones para el campo de selección.")
                return
        
        self.campo_creado = campo
        self.ventana.destroy()


def verificar_dependencias():
    try:
        from docx import Document
        return True
    except ImportError as e:
        print(f"""
        ❌ DEPENDENCIAS REQUERIDAS NO INSTALADAS
        
        Ejecute en la terminal:
        pip install python-docx
        
        Error: {e}
        """)
        return False

if __name__ == "__main__":
    if verificar_dependencias():
        app = SistemaPlantillasPersonalizadas()
        app.root.mainloop()