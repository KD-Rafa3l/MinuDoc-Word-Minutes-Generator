import re
from datetime import datetime

def validar_fecha_ddmmaaaa(texto):
    texto = texto.strip()
    patron = r"^(0[1-9]|[12][0-9]|3[01])/(0[1-9]|1[0-2])/(19|20)\d\d$"
    return bool(re.match(patron, texto))


import tkinter as tk
from tkinter import ttk, filedialog, messagebox, scrolledtext, simpledialog
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import re
from datetime import datetime
import json
from pathlib import Path

def aplicar_tema_moderno(root):
    style = ttk.Style(root)
    style.theme_use("clam")

    BG = "#f7f7f7"
    PANEL = "#ffffff"
    TEXT = "#474449"
    SUBTEXT = "#6b6b6b"
    BORDER = "#e0e0e0"
    BTN_BG = "#ffffff"
    BTN_BG_HOVER = "#f0f0f0"

    root.configure(bg=BG)

    style.configure(".", background=BG, foreground=TEXT, font=("Segoe UI", 10))
    style.configure("TFrame", background=BG)
    style.configure("TLabelFrame", background=BG, relief="solid", borderwidth=1, bordercolor=BORDER)
    style.configure("TLabelFrame.Label", background=BG, foreground=TEXT, font=("Segoe UI", 11, "bold"))
    style.configure("TLabel", background=BG, foreground=TEXT)
    style.configure("TEntry", relief="flat", borderwidth=1)
    style.configure("TCombobox", relief="flat", borderwidth=1)

    style.configure(
        "Notion.TButton",
        background=BTN_BG,
        foreground=TEXT,
        padding=10,
        borderwidth=1,
        relief="flat"
    )
    style.map(
        "Notion.TButton",
        background=[("active", BTN_BG_HOVER)],
        foreground=[("active", TEXT)]
    )

    style.configure("TNotebook", background=BG, borderwidth=0)
    style.configure("TNotebook.Tab", background=PANEL, foreground=TEXT, padding=[12, 8])
    style.map("TNotebook.Tab",
              background=[("selected", "#2563eb")],
              foreground=[("selected", "#ffffff")])


def aplicar_tema_oscuro(root):
    style = ttk.Style(root)
    style.theme_use("clam")

    BG = "#1e1e1e"
    PANEL = "#262626"
    TEXT = "#e5e5e5"
    SUB = "#a1a1a1"
    BORDER = "#333333"
    BTN_BG = "#2d2d2d"
    BTN_BG_HOVER = "#3a3a3a"

    root.configure(bg=BG)

    style.configure(".", background=BG, foreground=TEXT, font=("Segoe UI", 10))
    style.configure("TFrame", background=BG)
    style.configure("TLabelFrame", background=BG, relief="solid", borderwidth=1, bordercolor=BORDER)
    style.configure("TLabelFrame.Label", background=BG, foreground=TEXT, font=("Segoe UI", 11, "bold"))
    style.configure("TLabel", background=BG, foreground=TEXT)
    style.configure("TEntry", relief="flat", borderwidth=1, fieldbackground=PANEL, foreground=TEXT, insertcolor=TEXT)
    style.configure("TCombobox", relief="flat", borderwidth=1, fieldbackground=PANEL, background=PANEL, foreground=TEXT)

    style.configure(
        "Notion.TButton",
        background=BTN_BG,
        foreground=TEXT,
        padding=10,
        borderwidth=1,
        relief="flat"
    )
    style.map(
        "Notion.TButton",
        background=[("active", BTN_BG_HOVER)],
        foreground=[("active", TEXT)]
    )

    style.configure("TNotebook", background=BG, borderwidth=0)
    style.configure("TNotebook.Tab", background=PANEL, foreground=TEXT, padding=[12, 8])
    style.map("TNotebook.Tab",
              background=[("selected", BTN_BG)],
              foreground=[("selected", TEXT)])


class ScrollableFrame(ttk.Frame):
    def __init__(self, container, *args, **kwargs):
        super().__init__(container, *args, **kwargs)

        # Canvas y scrollbar vertical (pack para integrarse con padres que usan pack)
        self.canvas = tk.Canvas(self, highlightthickness=0)
        self.v_scrollbar = ttk.Scrollbar(self, orient="vertical", command=self.canvas.yview)
        self.h_scrollbar = ttk.Scrollbar(self, orient="horizontal", command=self.canvas.xview)

        # Layout: scrollbar derecha, canvas ocupa resto y scrollbar horizontal abajo
        self.v_scrollbar.pack(side="right", fill="y")
        self.canvas.pack(side="left", fill="both", expand=True)

        # Conectar scrollbars
        self.canvas.configure(yscrollcommand=self.v_scrollbar.set, xscrollcommand=self.h_scrollbar.set)

        # Frame interior donde se colocarán los widgets
        self.scrollable_frame = ttk.Frame(self.canvas)
        self.canvas_frame = self.canvas.create_window((0, 0), window=self.scrollable_frame, anchor="nw")

        # Actualizar scrollregion cuando cambie el contenido interior
        self.scrollable_frame.bind("<Configure>", lambda e: self.canvas.configure(scrollregion=self.canvas.bbox("all")))

        # Ajustar el ancho del frame interior al ancho del canvas (evita horizontales innecesarias)
        def _on_canvas_configure(event):
            try:
                self.canvas.itemconfig(self.canvas_frame, width=event.width)
                self.canvas.configure(scrollregion=self.canvas.bbox("all"))
            except Exception:
                pass
        self.canvas.bind("<Configure>", _on_canvas_configure)

        # Soporte rueda del ratón
        self._bind_mousewheel()

    def _on_mousewheel_windows(self, event):
        self.canvas.yview_scroll(int(-1 * (event.delta / 120)), "units")

    def _on_mousewheel_darwin(self, event):
        self.canvas.yview_scroll(int(-1 * event.delta), "units")

    def _on_mousewheel_linux_up(self, event):
        self.canvas.yview_scroll(-1, "units")

    def _on_mousewheel_linux_down(self, event):
        self.canvas.yview_scroll(1, "units")

    def _bind_mousewheel(self):
        # Vincular la rueda globalmente para que funcione cuando el cursor esté en la ventana
        self.canvas.bind_all("<MouseWheel>", lambda e: self._on_mousewheel_windows(e), add="+")
        self.canvas.bind_all("<Button-4>", lambda e: self._on_mousewheel_linux_up(e), add="+")
        self.canvas.bind_all("<Button-5>", lambda e: self._on_mousewheel_linux_down(e), add="+")


class SistemaPlantillasPersonalizadas:
    def __init__(self):
        self.root = tk.Tk()
        self.root.title("Sistema de Plantillas para Minutas Jurídicas - Versión Mejorada")
        self.root.geometry("1400x900")
        self.setup_icon()
        
        # Configurar ventana redimensionable
        self.root.resizable(True, True)
        self.root.minsize(1200, 700)
        
        # Variables de estado
        self.plantillas_personalizadas = {}
        self.plantilla_activa = None
        
        # Crear carpeta de plantillas
        self.carpeta_plantillas = Path("plantillas_personalizadas")
        self.carpeta_plantillas.mkdir(exist_ok=True)
        
        self.configurar_interfaz()
        self.cargar_plantillas_guardadas()
    
    def setup_icon(self):
        # Manejo seguro de icono

        try:
            self.root.iconbitmap("law_icon.ico")
        except:
            pass
    

    def configurar_interfaz(self):
        aplicar_tema_moderno(self.root)
        self.modo_oscuro = False

        cont = ttk.Frame(self.root, padding=40)
        cont.pack(fill="both", expand=True)

        header = ttk.Frame(cont)
        header.pack(fill="x", pady=(0, 25))

        titulo = ttk.Label(
            header,
            text="SISTEMA DE PLANTILLAS PARA MINUTAS JURÍDICAS",
            font=("Segoe UI", 22, "bold")
        )
        titulo.pack(anchor="w", pady=4)

        subtitulo = ttk.Label(
            header,
            text="Gestione, edite y genere minutas jurídicas de forma organizada y moderna.",
            font=("Segoe UI", 11)
        )
        subtitulo.pack(anchor="w")

        def toggle_dark():
            self.modo_oscuro = not self.modo_oscuro
            if self.modo_oscuro:
                aplicar_tema_oscuro(self.root)
                self.status_var.set("Modo oscuro activado")
            else:
                aplicar_tema_moderno(self.root)
                self.status_var.set("Modo claro activado")

        ttk.Button(
            header,
            text="Modo oscuro",
            style="Notion.TButton",
            command=toggle_dark
        ).pack(anchor="e", pady=5)

        quick = ttk.Frame(cont)
        quick.pack(fill="x", pady=(10, 25))

        fila1 = ttk.Frame(quick)
        fila1.pack(pady=5)

        ttk.Button(fila1, text="Crear nueva plantilla", style="Notion.TButton",
                   command=self.crear_plantilla_desde_minuta).pack(side="left", padx=6)
        ttk.Button(fila1, text="Editar plantilla", style="Notion.TButton",
                   command=self.editar_plantilla).pack(side="left", padx=6)
        ttk.Button(fila1, text="Generar minuta", style="Notion.TButton",
                   command=self.generar_minuta).pack(side="left", padx=6)

        fila2 = ttk.Frame(quick)
        fila2.pack(pady=5)

        ttk.Button(fila2, text="Importar", style="Notion.TButton",
                   command=self.importar_plantilla).pack(side="left", padx=6)
        ttk.Button(fila2, text="Exportar", style="Notion.TButton",
                   command=self.exportar_plantilla).pack(side="left", padx=6)
        ttk.Button(fila2, text="Limpiar", style="Notion.TButton",
                   command=self.limpiar_formulario).pack(side="left", padx=6)
        ttk.Button(fila2, text="Eliminar", style="Notion.TButton",
                   command=self.eliminar_plantilla_activa).pack(side="left", padx=6)

        active = ttk.LabelFrame(cont, text="Plantilla activa", padding=20)
        active.pack(fill="x", pady=(10, 25))

        filaA = ttk.Frame(active)
        filaA.pack(fill="x")

        ttk.Label(filaA, text="Plantilla actual:", font=("Segoe UI", 10, "bold")).grid(row=0, column=0, sticky="w")

        self.combo_plantillas = ttk.Combobox(filaA, width=35, state="readonly", font=("Segoe UI", 10))
        self.combo_plantillas.grid(row=0, column=1, padx=15)
        self.combo_plantillas.bind('<<ComboboxSelected>>', self.cambiar_plantilla)

        ttk.Label(filaA, text="Descripción:", font=("Segoe UI", 10)).grid(row=0, column=2, padx=(10, 5))
        self.label_descripcion = ttk.Label(filaA, text="Ninguna plantilla seleccionada")
        self.label_descripcion.grid(row=0, column=3, sticky="w")

        notebook_container = ttk.Frame(cont)
        notebook_container.pack(fill="both", expand=True)

        self.notebook = ttk.Notebook(notebook_container)
        self.notebook.pack(fill="both", expand=True)

        self.tab_formulario = ttk.Frame(self.notebook)
        self.notebook.add(self.tab_formulario, text="Formulario")
        self.configurar_tab_formulario()

        self.tab_vista_previa = ttk.Frame(self.notebook)
        self.notebook.add(self.tab_vista_previa, text="Vista previa")
        self.configurar_tab_vista_previa()

        self.tab_plantillas = ttk.Frame(self.notebook)
        self.notebook.add(self.tab_plantillas, text="Gestión")
        self.configurar_tab_plantillas()

        status_frame = ttk.Frame(cont)
        status_frame.pack(fill="x", pady=(20, 0))

        self.status_var = tk.StringVar(value="Sistema listo - Seleccione o cree una plantilla para comenzar")
        ttk.Label(status_frame, textvariable=self.status_var).pack(fill="x")

    def configurar_estilos(self):
        """Configurar estilos visuales"""
        style = ttk.Style()
        style.configure("TLabel", font=("Arial", 10))
        style.configure("TButton", font=("Arial", 10))
        style.configure("TEntry", font=("Arial", 10))
        style.configure("TCombobox", font=("Arial", 10))
    

    def configurar_tab_formulario(self):
        cont = ttk.Frame(self.tab_formulario, padding=25)
        cont.pack(fill="both", expand=True)

        info = ttk.LabelFrame(cont, text="Plantilla activa", padding=1)
        info.pack(fill="x", pady=(0, 10))

        # Cambiar a grid para el frame principal de info
        fila = ttk.Frame(info)
        fila.pack(fill="x")

        self.label_info_nombre = ttk.Label(fila, text="Nombre: -", font=("Segoe UI", 11, "bold"))
        self.label_info_nombre.pack(side="left", padx=(0, 10))

        self.label_info_campos = ttk.Label(fila, text="Campos: 0", font=("Segoe UI", 10))
        self.label_info_campos.pack(side="left", padx=(0, 60))

        self.label_info_desc = ttk.Label(fila, text="Descripción: -", font=("Segoe UI", 10))
        self.label_info_desc.pack(side="left")

        # Botones rápidos relacionados al formulario (generar Word / vista previa)
        acciones_form = ttk.Frame(fila)  # Cambiado: ahora es hijo de fila
        acciones_form.pack(side="right", padx=(10, 0))  # Usar pack en lugar de grid

        ttk.Button(
            acciones_form,
            text="🖨️ Generar Word",
            style="Notion.TButton",
            command=self.generar_minuta
        ).pack(side="left", padx=4)  # Usar pack en lugar de side="right"

        ttk.Button(
            acciones_form,
            text="👁️ Vista previa",
            style="Notion.TButton",
            command=lambda: self.notebook.select(1)
        ).pack(side="left", padx=4)  # Usar pack en lugar de side="right"

        form_panel = ttk.LabelFrame(cont, text="Campos", padding=10)
        form_panel.pack(fill="both", expand=True)

        # Forzar layout con grid para que "Campos" y su scrollbar ocupen TODO el espacio
        form_panel.grid_rowconfigure(0, weight=1)
        form_panel.grid_columnconfigure(0, weight=1)

        wrapper = ttk.Frame(form_panel)
        wrapper.grid(row=0, column=0, sticky="nsew")             # usar grid en el padre
        wrapper.grid_rowconfigure(0, weight=1)
        wrapper.grid_columnconfigure(0, weight=1)

        # Contenedor scrollable (canvas + scrollbar) dentro del wrapper
        self.frame_campos_container = ScrollableFrame(wrapper)
        self.frame_campos_container.grid(row=0, column=0, sticky="nsew")

        
        # Forzar recalculo del canvas cuando wrapper cambie de tamaño
        def _on_wrapper_configure(event):
            try:
                c = self.frame_campos_container.canvas
                c.update_idletasks()
                bbox = c.bbox("all")
                if bbox:
                    c.configure(scrollregion=bbox)
            except Exception:
                pass
        wrapper.bind("<Configure>", _on_wrapper_configure)
         
         # El frame donde realmente se colocan los widgets (usar en el resto del código)
        self.frame_campos = self.frame_campos_container.scrollable_frame

        self.label_form_vacio = ttk.Label(
            self.frame_campos,
            text="No hay plantilla activa.\nSeleccione una desde arriba.",
            font=("Segoe UI", 12),
            foreground="#7a7a7a",
            justify="center"
        )
        self.label_form_vacio.pack(pady=3)


    def configurar_tab_vista_previa(self):
        cont = ttk.Frame(self.tab_vista_previa, padding=25)
        cont.pack(fill="both", expand=True)

        title = ttk.Frame(cont)
        title.pack(fill="x", pady=(0, 20))

        ttk.Label(
            title,
            text="Vista previa de minuta",
            font=("Segoe UI", 16, "bold")
        ).pack(side="left")

        # Botón principal para generar la minuta (reemplaza marcadores y muestra vista previa)
        ttk.Button(
            title,
            text="Generar minuta",
            style="Notion.TButton",
            command=self.generar_minuta
        ).pack(side="right", padx=6)

        # Botón adicional: guarda directamente el contenido de la vista previa en .docx
        ttk.Button(
            title,
            text="Guardar Word (vista previa)",
            style="Notion.TButton",
            command=self.guardar_word_desde_vista
        ).pack(side="right", padx=6)

        prev = ttk.LabelFrame(cont, text="Contenido generado", padding=20)
        prev.pack(fill="both", expand=True)

        self.texto_vista_previa = scrolledtext.ScrolledText(
            prev,
            wrap="word",
            font=("Consolas", 11),
            background="#ffffff",
            padx=15,
            pady=15,
            relief="flat",
            borderwidth=1
        )
        self.texto_vista_previa.pack(fill="both", expand=True)

        self.texto_vista_previa.insert("1.0", "Aún no has generado la minuta.")


    def configurar_tab_plantillas(self):
        cont = ttk.Frame(self.tab_plantillas, padding=25)
        cont.pack(fill="both", expand=True)

        lista_panel = ttk.LabelFrame(cont, text="Plantillas disponibles", padding=20)
        lista_panel.pack(fill="both", expand=True, pady=(0, 25))

        top = ttk.Frame(lista_panel)
        top.pack(fill="x", pady=(0, 15))

        ttk.Label(top, text="Seleccione una plantilla").pack(side="left")

        listado = ttk.Frame(lista_panel)
        listado.pack(fill="both", expand=True)

        self.lista_plantillas = tk.Listbox(
            listado,
            height=10,
            font=("Segoe UI", 11),
            borderwidth=0,
            highlightthickness=1,
            relief="flat"
        )
        self.lista_plantillas.pack(side="left", fill="both", expand=True)

        scroll = ttk.Scrollbar(listado, orient="vertical", command=self.lista_plantillas.yview)
        scroll.pack(side="right", fill="y")
        self.lista_plantillas.configure(yscrollcommand=scroll.set)

        botones = ttk.Frame(lista_panel)
        botones.pack(fill="x", pady=15)

        ttk.Button(botones, text="Ver detalles", style="Notion.TButton",
                   command=self.ver_detalles_plantilla).pack(side="left", padx=8)
        ttk.Button(botones, text="Editar", style="Notion.TButton",
                   command=self.editar_plantilla).pack(side="left", padx=8)
        ttk.Button(botones, text="Probar", style="Notion.TButton",
                   command=self.probar_plantilla).pack(side="left", padx=8)

        detalles = ttk.LabelFrame(cont, text="Detalles", padding=20)
        detalles.pack(fill="x")

        self.texto_detalles = scrolledtext.ScrolledText(
            detalles,
            wrap="word",
            height=8,
            font=("Segoe UI", 10),
            background="#ffffff",
            relief="flat",
            borderwidth=1,
            padx=15,
            pady=15
        )
        self.texto_detalles.pack(fill="both", expand=True)

        self.texto_detalles.insert("1.0", "Seleccione una plantilla para ver sus detalles.")

    def crear_plantilla_desde_minuta(self):
        archivo = filedialog.askopenfilename(
            title="Seleccionar minuta base para crear plantilla",
            filetypes=[("Documentos Word", "*.docx"), ("Todos los archivos", "*.*")]
        )
        
        if archivo:
            try:
                doc = Document(archivo)
                contenido = ""
                
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        contenido += paragraph.text + "\n\n"
                
                editor = EditorPlantillasDesdeMinuta(self.root, self.carpeta_plantillas, contenido, archivo)
                self.root.wait_window(editor.ventana)
                self.cargar_plantillas_guardadas()
                
            except Exception as e:
                messagebox.showerror("Error", f"No se pudo cargar la minuta: {str(e)}")
    
    def editar_plantilla(self):
        # Preferir la selección en la lista; si no hay, usar la plantilla activa en el combobox
        nombre_plantilla = None
        seleccion = self.lista_plantillas.curselection()
        if seleccion:
            nombre_plantilla = self.lista_plantillas.get(seleccion[0])
        else:
            nombre_combo = self.combo_plantillas.get()
            if nombre_combo in self.plantillas_personalizadas:
                nombre_plantilla = nombre_combo

        if not nombre_plantilla:
            messagebox.showwarning("Advertencia", "Seleccione una plantilla (lista o combobox) para editar.")
            return

        plantilla = self.plantillas_personalizadas.get(nombre_plantilla)
        if plantilla:
            editor = EditorPlantillasDesdeMinuta(
                self.root, self.carpeta_plantillas,
                plantilla.get('contenido_base', ''),
                plantilla.get('documento_origen', ''),
                plantilla_existente=plantilla
            )
            self.root.wait_window(editor.ventana)
            # Recargar plantillas y reactivar la misma plantilla si existe
            self.cargar_plantillas_guardadas()
            if nombre_plantilla in self.plantillas_personalizadas:
                self.combo_plantillas.set(nombre_plantilla)
                self.cambiar_plantilla()
        else:
            messagebox.showwarning("Advertencia", "Seleccione una plantilla de la lista para editar.")
    
    def generar_minuta(self):
        if not self.plantilla_activa:
            messagebox.showwarning("Advertencia", "No hay plantilla activa. Seleccione una plantilla primero.")
            return
        
        datos = self.obtener_datos_formulario()
        errores = self.validar_formulario(datos)
        
        if errores:
            messagebox.showwarning("Campos requeridos", 
                                "Los siguientes campos son requeridos:\n\n• " + "\n• ".join(errores))
            return
        
        try:
            minuta_generada = self.aplicar_plantilla(self.plantilla_activa, datos)
            
            self.texto_vista_previa.delete("1.0", tk.END)
            self.texto_vista_previa.insert("1.0", minuta_generada)
            
            self.generar_documento_word(minuta_generada)
            
            self.notebook.select(1)
            self.status_var.set("✅ Minuta generada y guardada exitosamente!")
            
        except Exception as e:
            messagebox.showerror("Error", f"No se pudo generar la minuta: {str(e)}")
    
    def obtener_datos_formulario(self):
        datos = {}
        if hasattr(self, 'campos_ui'):
            for campo_id, widget_info in self.campos_ui.items():
                widget = widget_info['widget']
                if isinstance(widget, ttk.Entry):
                    datos[campo_id] = widget.get()
                elif isinstance(widget, tk.Text):
                    datos[campo_id] = widget.get("1.0", tk.END).strip()
                elif isinstance(widget, ttk.Combobox):
                    datos[campo_id] = widget.get()
        return datos
    
    def validar_formulario(self, datos):
        errores = []
        if hasattr(self, 'campos_ui'):
            for campo_id, widget_info in self.campos_ui.items():
                if widget_info.get('requerido', False) and not datos.get(campo_id):
                    errores.append(widget_info['label'])
        return errores
    
    def aplicar_plantilla(self, plantilla, datos):
        contenido_base = plantilla.get('contenido_base', '')
        
        for campo_id, valor in datos.items():
            marcador = f"[[{campo_id}]]"
            contenido_base = contenido_base.replace(marcador, valor)
        
        contenido_base = re.sub(r'\[\[.*?\]\]', '[SIN DATO]', contenido_base)
        return contenido_base
    
    def generar_documento_word(self, contenido):
        doc = Document()
        self.aplicar_formato_apa(doc)
        
        for linea in contenido.split('\n'):
            if linea.strip():
                doc.add_paragraph(linea)
        
        archivo_salida = filedialog.asksaveasfilename(
            title="Guardar minuta como...",
            defaultextension=".docx",
            filetypes=[("Documentos Word", "*.docx")],
            initialfile=f"minuta_{datetime.now().strftime('%Y%m%d_%H%M')}.docx"
        )
        
        if archivo_salida:
            doc.save(archivo_salida)
            try:
                os.startfile(archivo_salida)
            except Exception:
                pass  # Evita crash si el SO no soporta startfile
            return True
        return False
    
    def guardar_word_desde_vista(self):
        """Guardar el contenido actual de la vista previa en un .docx."""
        contenido = self.texto_vista_previa.get("1.0", tk.END).strip()
        if not contenido:
            messagebox.showwarning("Advertencia", "No hay contenido en la vista previa para guardar.")
            return
        self.generar_documento_word(contenido)
    
    def aplicar_formato_apa(self, doc):
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)
        
        paragraph_format = style.paragraph_format
        paragraph_format.line_spacing = 2.0
    
    def cargar_plantillas_guardadas(self):
        self.plantillas_personalizadas = {}
        for archivo in self.carpeta_plantillas.glob("*.json"):
            try:
                with open(archivo, 'r', encoding='utf-8') as f:
                    plantilla = json.load(f)
                self.plantillas_personalizadas[archivo.stem] = plantilla
            except Exception as e:
                print(f"Error cargando plantilla {archivo}: {e}")
        
        self.actualizar_listas_plantillas()
    
    def actualizar_listas_plantillas(self):
        plantillas = list(self.plantillas_personalizadas.keys())
        self.combo_plantillas['values'] = plantillas
        self.lista_plantillas.delete(0, tk.END)
        
        for nombre in plantillas:
            self.lista_plantillas.insert(tk.END, nombre)
        
        if plantillas:
            self.combo_plantillas.set(plantillas[0])
            self.cambiar_plantilla()
    
    def cambiar_plantilla(self, event=None):
        nombre_plantilla = self.combo_plantillas.get()
        if nombre_plantilla in self.plantillas_personalizadas:
            self.plantilla_activa = self.plantillas_personalizadas[nombre_plantilla]
            # DEBUG: mostrar en consola cuántos campos trae la plantilla activa
            campos = self.plantilla_activa.get('campos_personalizados', [])
            print(f"[DEBUG] cambiar_plantilla -> '{nombre_plantilla}' cargada. campos_personalizados: {len(campos)}")
            # Siempre mostrar el formulario cuando se cambie de plantilla
            self.cargar_formulario_plantilla()
            self.actualizar_info_plantilla()
            if hasattr(self, "notebook"):
                self.notebook.select(0)
            self.status_var.set(f"✅ Plantilla activa: {nombre_plantilla}")
    
    def cargar_formulario_plantilla(self):
        # DEBUG: confirmar entrada a la función y contenido de plantilla_activa
        print(f"[DEBUG] cargar_formulario_plantilla llamado. plantilla_activa presente: {self.plantilla_activa is not None}")
        if self.plantilla_activa is not None:
            print(f"[DEBUG] campos_personalizados (preview): {self.plantilla_activa.get('campos_personalizados', [])[:3]}")
        for widget in self.frame_campos.winfo_children():
            widget.destroy()
        
        self.campos_ui = {}
        
        if not self.plantilla_activa:
            # Forzar actualización del scrollregion (vacío)
            if hasattr(self, 'frame_campos_container'):
                try:
                    c = self.frame_campos_container.canvas
                    c.update_idletasks()
                    c.configure(scrollregion=c.bbox("all"))
                    c.yview_moveto(0.0)
                except Exception:
                    pass
            return
        
        campos = self.plantilla_activa.get('campos_personalizados', [])
        
        if not campos:
            self.label_form_vacio = ttk.Label(self.frame_campos, 
                                             text="Esta plantilla no tiene campos personalizados definidos",
                                             font=("Arial", 11), foreground="gray")
            self.label_form_vacio.pack(pady=50)
            # Actualizar scrollregion cuando hay sólo el mensaje vacío
            if hasattr(self, 'frame_campos_container'):
                try:
                    c = self.frame_campos_container.canvas
                    c.update_idletasks()
                    c.configure(scrollregion=c.bbox("all"))
                    c.yview_moveto(0.0)
                except Exception:
                    pass
            return
        
        for i, campo in enumerate(campos):
            self.crear_campo_formulario(campo, i)

        # Después de añadir todos los campos, forzar recálculo del scrollregion
        if hasattr(self, 'frame_campos_container'):
            try:
                c = self.frame_campos_container.canvas
                # asegurar que se actualice layout antes de calcular bbox
                c.update_idletasks()
                bbox = c.bbox("all")
                if bbox:
                    c.configure(scrollregion=bbox)
                c.yview_moveto(0.0)  # colocar al inicio
            except Exception as e:
                print(f"[DEBUG] error actualizando scrollregion: {e}")
    
    def crear_campo_formulario(self, campo, index):
        frame_campo = ttk.Frame(self.frame_campos)
        frame_campo.pack(fill="x", pady=8, padx=15)
        
        label_text = campo['nombre']
        if campo.get('requerido', False):
            label_text += " *"
        
        label = ttk.Label(frame_campo, text=label_text, width=25, anchor="w", font=("Arial", 10))
        label.pack(side="left", padx=(0, 15))
        
        campo_id = campo['id']
        tipo = campo.get('tipo', 'texto')
        if tipo == 'texto':
            widget = ttk.Entry(frame_campo, width=50, font=("Arial", 9))
            widget.pack(side="left", fill="x", expand=True)

        elif tipo == 'textarea':
            frame_text = ttk.Frame(frame_campo)
            frame_text.pack(side="left", fill="x", expand=True)

            widget = tk.Text(frame_text, width=60, height=4, wrap=tk.WORD, font=("Arial", 9))
            scrollbar = ttk.Scrollbar(frame_text, orient="vertical", command=widget.yview)
            widget.configure(yscrollcommand=scrollbar.set)

            widget.pack(side="left", fill="both", expand=True)
            scrollbar.pack(side="right", fill="y")

        elif tipo == 'seleccion':
            widget = ttk.Combobox(frame_campo, width=48, values=campo.get('opciones', []), font=("Arial", 9))
            widget.pack(side="left", fill="x", expand=True)

        elif tipo == 'fecha':
            widget = ttk.Entry(frame_campo, width=25, font=("Arial", 9))
            widget.pack(side="left")
            ttk.Label(frame_campo, text="(DD/MM/AAAA)", font=("Arial", 8), foreground="gray").pack(side="left", padx=(5, 0))
        else:
            # Tipo de campo no reconocido: usar una entrada de texto básica
            widget = ttk.Entry(frame_campo, width=50, font=("Arial", 9))
            widget.pack(side="left", fill="x", expand=True)

        if campo.get('descripcion'):
            self.crear_tooltip(label, campo['descripcion'])

        self.campos_ui[campo_id] = {
            'widget': widget,
            'label': campo['nombre'],
            'requerido': campo.get('requerido', False)
        }

    def crear_tooltip(self, widget, text):
        def on_enter(event):
            tooltip = tk.Toplevel()
            tooltip.wm_overrideredirect(True)
            tooltip.wm_geometry(f"+{event.x_root+10}+{event.y_root+10}")
            label = ttk.Label(tooltip, text=text, background="lightyellow", 
                            relief="solid", borderwidth=1, padding=5, font=("Arial", 9))
            label.pack()
            widget.tooltip = tooltip
        
        def on_leave(event):
            if hasattr(widget, 'tooltip'):
                widget.tooltip.destroy()
        
        widget.bind("<Enter>", on_enter)
        widget.bind("<Leave>", on_leave)
    
    def actualizar_info_plantilla(self):
        if self.plantilla_activa:
            self.label_info_nombre.config(text=f"Nombre: {self.plantilla_activa.get('nombre', 'N/A')}")
            self.label_info_desc.config(text=f"Descripción: {self.plantilla_activa.get('descripcion', 'N/A')}")
            
            campos = self.plantilla_activa.get('campos_personalizados', [])
            campos_requeridos = sum(1 for c in campos if c.get('requerido', False))
            self.label_info_campos.config(text=f"Campos: {len(campos)} (Requeridos: {campos_requeridos})")
            
            self.label_descripcion.config(text=self.plantilla_activa.get('descripcion', 'Sin descripción'))
    
    def importar_plantilla(self):
        archivo = filedialog.askopenfilename(
            title="Importar plantilla",
            filetypes=[("Archivos de plantilla", "*.json")]
        )
        
        if archivo:
            try:
                with open(archivo, 'r', encoding='utf-8') as f:
                    plantilla = json.load(f)
                
                nombre = plantilla.get('nombre', Path(archivo).stem)
                
                if nombre in self.plantillas_personalizadas:
                    respuesta = messagebox.askyesno("Confirmar", 
                                                  f"¿Sobrescribir la plantilla existente '{nombre}'?")
                    if not respuesta:
                        return
                
                archivo_destino = self.carpeta_plantillas / f"{nombre}.json"
                with open(archivo_destino, 'w', encoding='utf-8') as f:
                    json.dump(plantilla, f, ensure_ascii=False, indent=2)
                
                self.cargar_plantillas_guardadas()
                messagebox.showinfo("Éxito", f"Plantilla '{nombre}' importada correctamente.")
                
            except Exception as e:
                messagebox.showerror("Error", f"No se pudo importar la plantilla: {str(e)}")
    
    def exportar_plantilla(self):
        seleccion = self.lista_plantillas.curselection()
        if seleccion:
            nombre_plantilla = self.lista_plantillas.get(seleccion[0])
            plantilla = self.plantillas_personalizadas.get(nombre_plantilla)
            
            if plantilla:
                archivo = filedialog.asksaveasfilename(
                    title="Exportar plantilla como...",
                    defaultextension=".json",
                    filetypes=[("Archivos de plantilla", "*.json")],
                    initialfile=f"{nombre_plantilla}.json"
                )
                
                if archivo:
                    try:
                        with open(archivo, 'w', encoding='utf-8') as f:
                            json.dump(plantilla, f, ensure_ascii=False, indent=2)
                        messagebox.showinfo("Éxito", f"Plantilla exportada a: {archivo}")
                    except Exception as e:
                        messagebox.showerror("Error", f"No se pudo exportar: {str(e)}")
        else:
            messagebox.showwarning("Advertencia", "Seleccione una plantilla para exportar.")
    
    def eliminar_plantilla_activa(self):
        nombre_plantilla = self.combo_plantillas.get()
        if nombre_plantilla and nombre_plantilla in self.plantillas_personalizadas:
            respuesta = messagebox.askyesno("Confirmar", 
                                          f"¿Está seguro de eliminar la plantilla '{nombre_plantilla}'?")
            if respuesta:
                archivo_plantilla = self.carpeta_plantillas / f"{nombre_plantilla}.json"
                if archivo_plantilla.exists():
                    archivo_plantilla.unlink()
                
                self.cargar_plantillas_guardadas()
                messagebox.showinfo("Éxito", f"Plantilla '{nombre_plantilla}' eliminada.")
        else:
            messagebox.showwarning("Advertencia", "No hay plantilla seleccionada para eliminar.")
    
    def probar_plantilla(self):
        seleccion = self.lista_plantillas.curselection()
        if seleccion:
            nombre_plantilla = self.lista_plantillas.get(seleccion[0])
            self.combo_plantillas.set(nombre_plantilla)
            self.cambiar_plantilla()
            self.notebook.select(0)
            messagebox.showinfo("Éxito", f"Plantilla '{nombre_plantilla}' activada para prueba.")
        else:
            messagebox.showwarning("Advertencia", "Seleccione una plantilla de la lista.")
    
    def ver_detalles_plantilla(self):
        seleccion = self.lista_plantillas.curselection()
        if seleccion:
            nombre_plantilla = self.lista_plantillas.get(seleccion[0])
            plantilla = self.plantillas_personalizadas.get(nombre_plantilla)
            
            if plantilla:
                detalles = f"""INFORMACIÓN DETALLADA DE LA PLANTILLA

Nombre: {plantilla.get('nombre', 'N/A')}
Descripción: {plantilla.get('descripcion', 'N/A')}
Tipo: {plantilla.get('tipo', 'N/A')}
Fecha creación: {plantilla.get('fecha_creacion', 'N/A')}
Documento origen: {plantilla.get('documento_origen', 'N/A')}

CAMPOS PERSONALIZADOS:
"""
            campos = plantilla.get('campos_personalizados', [])
            for i, campo in enumerate(campos, 1):
                requerido = "SÍ" if campo.get('requerido') else "no"
                detalles += f"\n{i}. {campo['nombre']} ({campo['tipo']}) - Requerido: {requerido}"
                if campo.get('descripcion'):
                    detalles += f"\n   Descripción: {campo['descripcion']}"
            
            self.texto_detalles.delete("1.0", tk.END)
            self.texto_detalles.insert("1.0", detalles)
        else:
            messagebox.showwarning("Advertencia", "Seleccione una plantilla de la lista.")
    
    def limpiar_formulario(self):
        if hasattr(self, 'campos_ui'):
            for campo_id, widget_info in self.campos_ui.items():
                widget = widget_info['widget']
                if isinstance(widget, ttk.Entry) or isinstance(widget, ttk.Combobox):
                    widget.delete(0, tk.END)
                elif isinstance(widget, tk.Text):
                    widget.delete("1.0", tk.END)
        
        self.texto_vista_previa.delete("1.0", tk.END)
        self.texto_vista_previa.insert(tk.END, "Formulario limpiado. Complete los campos y genere una nueva minuta.")
        self.status_var.set("Formulario limpiado - Listo para nuevo proceso")


class EditorPlantillasDesdeMinuta:
    def __init__(self, parent, carpeta_plantillas, contenido_minuta="", archivo_origen="", plantilla_existente=None):
        self.parent = parent
        self.carpeta_plantillas = carpeta_plantillas
        self.contenido_minuta = contenido_minuta
        self.archivo_origen = archivo_origen
        self.plantilla_existente = plantilla_existente
        
        self.ventana = tk.Toplevel(parent)
        self.ventana.title("Editor de Plantillas - Crear/Editar Plantilla")
        self.ventana.geometry("1400x900")
        self.ventana.transient(parent)
        self.ventana.grab_set()
        self.ventana.resizable(True, True)
        self.ventana.minsize(1200, 700)
        
        self.campos_personalizados = []
        self.mapeo_selecciones = {}
        self.texto_seleccionado_actual = None
        self.posicion_seleccion_actual = None
        
        # Frame principal con scroll
        self.main_scrollable = ScrollableFrame(self.ventana)
        self.main_scrollable.pack(fill="both", expand=True)
        
        self.configurar_interfaz()
        
        if plantilla_existente:
            self.cargar_plantilla_existente(plantilla_existente)
    

    def configurar_interfaz(self):
        cont = ttk.Frame(self.main_scrollable.scrollable_frame, padding=35)
        cont.pack(fill="both", expand=True)

        header = ttk.Frame(cont)
        header.pack(fill="x", pady=(0, 25))

        ttk.Label(
            header,
            text="Editor de Plantillas",
            font=("Segoe UI", 22, "bold")
        ).pack(anchor="w")

        ttk.Label(
            header,
            text="Seleccione texto y conviértalo en campos personalizados.",
            font=("Segoe UI", 11)
        ).pack(anchor="w", pady=(4, 0))

        info = ttk.LabelFrame(cont, text="Información de plantilla", padding=20)
        info.pack(fill="x", pady=(0, 25))

        grid = ttk.Frame(info)
        grid.pack(fill="x")

        ttk.Label(grid, text="Nombre:", font=("Segoe UI", 10, "bold")).grid(row=0, column=0, sticky="w")
        self.entry_nombre = ttk.Entry(grid, width=50)
        self.entry_nombre.grid(row=0, column=1, padx=15, pady=6)

        ttk.Label(grid, text="Descripción:", font=("Segoe UI", 10, "bold")).grid(row=1, column=0, sticky="w")
        self.entry_descripcion = ttk.Entry(grid, width=50)
        self.entry_descripcion.grid(row=1, column=1, padx=15, pady=6)

        ttk.Label(grid, text="Tipo:", font=("Segoe UI", 10, "bold")).grid(row=2, column=0, sticky="w")
        self.combo_tipo = ttk.Combobox(grid, width=48,
                                       values=["Amparo", "Contrato", "Demanda", "Recurso", "General", "Solicitud"])
        self.combo_tipo.grid(row=2, column=1, padx=15, pady=6)
        self.combo_tipo.set("General")

        workspace = ttk.Panedwindow(cont, orient=tk.HORIZONTAL)
        workspace.pack(fill="both", expand=True)

        left = ttk.LabelFrame(workspace, text="Minuta base", padding=20)
        right = ttk.LabelFrame(workspace, text="Campos personalizados", padding=20)

        # Agregar paneles al PanedWindow (arrastrables)
        workspace.add(left, weight=3)
        workspace.add(right, weight=1)

        self.texto_minuta = scrolledtext.ScrolledText(
            left,
            wrap="word",
            font=("Consolas", 11),
            background="#ffffff",
            relief="flat",
            borderwidth=1,
            padx=15,
            pady=15
        )
        self.texto_minuta.pack(fill="both", expand=True)

        if self.contenido_minuta:
            self.texto_minuta.insert("1.0", self.contenido_minuta)

        # Tag para marcadores persistentes y tag para selección activa (temporal)
        self.texto_minuta.tag_configure("seleccionado", background="lightgreen", foreground="darkgreen")
        self.texto_minuta.tag_configure("sel_activa", background="#e6f7d9")
        # Mantener registro de selección y mostrarla visualmente mientras el usuario la mantiene
        self.texto_minuta.bind("<ButtonRelease-1>", self.guardar_seleccion_actual)
        self.texto_minuta.bind("<KeyRelease>", self.guardar_seleccion_actual)

        ttk.Button(right, text="Crear campo desde selección",
                   style="Notion.TButton", command=self.crear_campo_desde_seleccion).pack(fill="x", pady=6)
        ttk.Button(right, text="Agregar campo manual",
                   style="Notion.TButton", command=self.agregar_campo_manual).pack(fill="x", pady=6)

        lista = ttk.Frame(right)
        lista.pack(fill="both", expand=True, pady=(15, 0))

        self.lista_campos = tk.Listbox(
            lista,
            height=10,
            font=("Segoe UI", 10),
            highlightthickness=1,
            relief="flat"
        )
        self.lista_campos.pack(side="left", fill="both", expand=True)

        scroll = ttk.Scrollbar(lista, orient="vertical", command=self.lista_campos.yview)
        scroll.pack(side="right", fill="y")
        self.lista_campos.configure(yscrollcommand=scroll.set)

        botones = ttk.Frame(right)
        botones.pack(fill="x", pady=10)

        ttk.Button(botones, text="Editar", style="Notion.TButton",
                   command=self.editar_campo).pack(side="left", padx=5)
        ttk.Button(botones, text="Eliminar", style="Notion.TButton",
                   command=self.eliminar_campo).pack(side="left", padx=5)

        final = ttk.Frame(cont)
        final.pack(fill="x", pady=25)

        ttk.Button(final, text="Guardar plantilla",
                   style="Notion.TButton", command=self.guardar_plantilla).pack(side="left", padx=8)
        ttk.Button(final, text="Vista previa de marcadores",
                   style="Notion.TButton", command=self.mostrar_vista_previa).pack(side="left", padx=8)
        ttk.Button(final, text="Cerrar",
                   style="Notion.TButton", command=self.ventana.destroy).pack(side="left", padx=8)

    def guardar_seleccion_actual(self, event=None):
        try:
            # Obtener rangos reales de la selección y almacenarlos como índices de texto
            if self.texto_minuta.tag_ranges(tk.SEL):
                inicio = self.texto_minuta.index(tk.SEL_FIRST)
                fin = self.texto_minuta.index(tk.SEL_LAST)
                self.texto_seleccionado_actual = self.texto_minuta.get(inicio, fin)
                self.posicion_seleccion_actual = (inicio, fin)
            else:
                self.texto_seleccionado_actual = None
                self.posicion_seleccion_actual = None
        except Exception:
            self.texto_seleccionado_actual = None
            self.posicion_seleccion_actual = None
    
    def crear_campo_desde_seleccion(self):
        # Si no había seleccionado previamente, intentar leer la selección actual del widget
        if not self.texto_seleccionado_actual:
            try:
                if self.texto_minuta.tag_ranges(tk.SEL):
                    inicio = self.texto_minuta.index(tk.SEL_FIRST)
                    fin = self.texto_minuta.index(tk.SEL_LAST)
                    self.texto_seleccionado_actual = self.texto_minuta.get(inicio, fin)
                    self.posicion_seleccion_actual = (inicio, fin)
            except Exception:
                self.texto_seleccionado_actual = None

        if not self.texto_seleccionado_actual:
            messagebox.showwarning("Advertencia", "Primero seleccione texto en la minuta.")
            return
        
        texto_seleccionado = self.texto_seleccionado_actual
        dialogo = DialogoCampoDesdeSeleccion(self.ventana, texto_seleccionado)
        self.ventana.wait_window(dialogo.ventana)
        
        if dialogo.campo_creado:
            campo = dialogo.campo_creado
            self.campos_personalizados.append(campo)
            
            if self.posicion_seleccion_actual:
                inicio, fin = self.posicion_seleccion_actual
                marcador = f"[[{campo['id']}]]"
                
                self.texto_minuta.delete(inicio, fin)
                self.texto_minuta.insert(inicio, marcador)
                
                nuevo_fin = self.texto_minuta.index(f"{inicio} + {len(marcador)}c")
                self.texto_minuta.tag_add("seleccionado", inicio, nuevo_fin)
            
            self.mapeo_selecciones[campo['id']] = {
                'texto_original': texto_seleccionado,
                'marcador': marcador
            }
            
            self.actualizar_lista_campos()
            messagebox.showinfo("Éxito", f"Campo '{campo['nombre']}' creado correctamente.")
            
            self.texto_seleccionado_actual = None
            self.posicion_seleccion_actual = None
    
    def agregar_campo_manual(self):
        dialogo = DialogoCampoDesdeSeleccion(self.ventana, "")
        self.ventana.wait_window(dialogo.ventana)
        
        if dialogo.campo_creado:
            self.campos_personalizados.append(dialogo.campo_creado)
            self.actualizar_lista_campos()
            messagebox.showinfo("Éxito", f"Campo '{dialogo.campo_creado['nombre']}' agregado manualmente.")
    
    def editar_campo(self):
        seleccion = self.lista_campos.curselection()
        if not seleccion:
            messagebox.showwarning("Advertencia", "Seleccione un campo para editar.")
            return
        
        index = seleccion[0]
        campo_existente = self.campos_personalizados[index]
        
        dialogo = DialogoCampoDesdeSeleccion(self.ventana, "", campo_existente)
        self.ventana.wait_window(dialogo.ventana)
        
        if dialogo.campo_creado:
            self.campos_personalizados[index] = dialogo.campo_creado
            self.actualizar_lista_campos()
            messagebox.showinfo("Éxito", f"Campo '{dialogo.campo_creado['nombre']}' actualizado.")
    
    def eliminar_campo(self):
        seleccion = self.lista_campos.curselection()
        if not seleccion:
            messagebox.showwarning("Advertencia", "Seleccione un campo para eliminar.")
            return
        
        index = seleccion[0]
        campo = self.campos_personalizados[index]
        
        respuesta = messagebox.askyesno("Confirmar", 
                                      f"¿Está seguro de eliminar el campo '{campo['nombre']}'?")
        if respuesta:
            if campo['id'] in self.mapeo_selecciones:
                contenido_actual = self.texto_minuta.get("1.0", tk.END)
                marcador = f"[[{campo['id']}]]"
                texto_original = self.mapeo_selecciones[campo['id']]['texto_original']
                nuevo_contenido = contenido_actual.replace(marcador, texto_original)
                
                self.texto_minuta.delete("1.0", tk.END)
                self.texto_minuta.insert("1.0", nuevo_contenido)
                
                del self.mapeo_selecciones[campo['id']]
            
            self.campos_personalizados.pop(index)
            self.actualizar_lista_campos()
            messagebox.showinfo("Éxito", f"Campo '{campo['nombre']}' eliminado.")
    
    def actualizar_lista_campos(self):
        self.lista_campos.delete(0, tk.END)
        for campo in self.campos_personalizados:
            requerido = " *" if campo.get('requerido') else ""
            self.lista_campos.insert(tk.END, f"{campo['nombre']}{requerido} ({campo['tipo']})")
    
    def cargar_plantilla_existente(self, plantilla):
        self.entry_nombre.delete(0, tk.END)
        self.entry_nombre.insert(0, plantilla.get('nombre', ''))
        
        self.entry_descripcion.delete(0, tk.END)
        self.entry_descripcion.insert(0, plantilla.get('descripcion', ''))
        
        self.combo_tipo.set(plantilla.get('tipo', 'General'))
        
        self.texto_minuta.delete("1.0", tk.END)
        self.texto_minuta.insert("1.0", plantilla.get('contenido_base', ''))
        
        self.campos_personalizados = plantilla.get('campos_personalizados', [])
        self.actualizar_lista_campos()
        
        self.resaltar_marcadores()
    
    def resaltar_marcadores(self):
        contenido = self.texto_minuta.get("1.0", tk.END)
        for marcador in re.findall(r'\[\[.*?\]\]', contenido):
            inicio = "1.0"
            while True:
                inicio = self.texto_minuta.search(marcador, inicio, tk.END)
                if not inicio:
                    break
                fin = f"{inicio} + {len(marcador)}c"
                self.texto_minuta.tag_add("seleccionado", inicio, fin)
                inicio = fin
    
    def mostrar_vista_previa(self):
        if not self.campos_personalizados:
            messagebox.showinfo("Marcadores", "No hay campos creados todavía.")
            return
        
        marcadores = "\n".join([f"[[{campo['id']}]] - {campo['nombre']} ({campo['tipo']})" 
                              for campo in self.campos_personalizados])
        messagebox.showinfo("Marcadores Disponibles", 
                          f"Puede usar estos marcadores en el contenido:\n\n{marcadores}")
    
    def guardar_plantilla(self):
        nombre = self.entry_nombre.get().strip()
        descripcion = self.entry_descripcion.get().strip()
        tipo = self.combo_tipo.get()
        contenido = self.texto_minuta.get("1.0", tk.END).strip()
        
        if not nombre:
            messagebox.showwarning("Advertencia", "El nombre de la plantilla es requerido.")
            return
        
        if not self.campos_personalizados:
            messagebox.showwarning("Advertencia", "Debe crear al menos un campo para la plantilla.")
            return
        
        plantilla = {
            'nombre': nombre,
            'descripcion': descripcion,
            'tipo': tipo,
            'fecha_creacion': datetime.now().isoformat(),
            'campos_personalizados': self.campos_personalizados,
            'contenido_base': contenido,
            'documento_origen': self.archivo_origen
        }
        
        archivo_plantilla = self.carpeta_plantillas / f"{nombre}.json"
        
        try:
            with open(archivo_plantilla, 'w', encoding='utf-8') as f:
                json.dump(plantilla, f, ensure_ascii=False, indent=2)
            
            messagebox.showinfo("Éxito", f"Plantilla '{nombre}' guardada correctamente!")
            self.ventana.destroy()
            
        except Exception as e:
            messagebox.showerror("Error", f"No se pudo guardar la plantilla: {str(e)}")


class DialogoCampoDesdeSeleccion:
    def __init__(self, parent, texto_seleccionado="", campo_existente=None):
        self.parent = parent
        self.texto_seleccionado = texto_seleccionado
        self.campo_existente = campo_existente
        self.campo_creado = None
        
        self.ventana = tk.Toplevel(parent)
        self.ventana.title("Configurar Campo Personalizado")
        self.ventana.geometry("600x650")
        self.ventana.transient(parent)
        self.ventana.grab_set()
        self.ventana.resizable(True, True)
        self.ventana.minsize(550, 600)
        
        # Frame con scroll para el diálogo
        self.dialog_scrollable = ScrollableFrame(self.ventana)
        self.dialog_scrollable.pack(fill="both", expand=True)
        
        self.configurar_interfaz()
        
        if campo_existente:
            self.cargar_datos_existentes(campo_existente)
    

    def configurar_interfaz(self):
        cont = ttk.Frame(self.dialog_scrollable.scrollable_frame, padding=30)
        cont.pack(fill="both", expand=True)

        ttk.Label(
            cont,
            text="Configurar Campo",
            font=("Segoe UI", 18, "bold")
        ).pack(anchor="w", pady=(0, 20))

        if self.texto_seleccionado:
            marco = ttk.LabelFrame(cont, text="Texto seleccionado", padding=15)
            marco.pack(fill="x", pady=10)

            ttk.Label(
                marco,
                text=self.texto_seleccionado,
                background="#ffffff",
                wraplength=500,
                font=("Segoe UI", 10)
            ).pack(fill="x")

        panel = ttk.LabelFrame(cont, text="Propiedades del campo", padding=20)
        panel.pack(fill="both", expand=True, pady=(10, 20))

        grid = ttk.Frame(panel)
        grid.pack(fill="x")

        ttk.Label(grid, text="ID:", font=("Segoe UI", 10, "bold")).grid(row=0, column=0, sticky="w", pady=5)
        self.entry_id = ttk.Entry(grid, width=40)
        self.entry_id.grid(row=0, column=1, pady=5, padx=10)

        ttk.Label(grid, text="Nombre visible:", font=("Segoe UI", 10, "bold")).grid(row=1, column=0, sticky="w", pady=5)
        self.entry_nombre = ttk.Entry(grid, width=40)
        self.entry_nombre.grid(row=1, column=1, pady=5, padx=10)

        ttk.Label(grid, text="Tipo:", font=("Segoe UI", 10, "bold")).grid(row=2, column=0, sticky="nw", pady=5)

        tipos = ttk.Frame(grid)
        tipos.grid(row=2, column=1, sticky="w", pady=5, padx=10)

        self.tipo_var = tk.StringVar(value="texto")

        tk.Radiobutton(tipos, text="Texto corto", variable=self.tipo_var, value="texto",
                       font=("Segoe UI", 10)).pack(anchor="w")
        tk.Radiobutton(tipos, text="Texto largo", variable=self.tipo_var, value="textarea",
                       font=("Segoe UI", 10)).pack(anchor="w")
        tk.Radiobutton(tipos, text="Selección", variable=self.tipo_var, value="seleccion",
                       font=("Segoe UI", 10)).pack(anchor="w")
        tk.Radiobutton(tipos, text="Fecha", variable=self.tipo_var, value="fecha",
                       font=("Segoe UI", 10)).pack(anchor="w")

        self.frame_opciones = ttk.LabelFrame(panel, text="Opciones (para selección)", padding=15)
        self.frame_opciones.pack(fill="x", pady=15)

        self.texto_opciones = tk.Text(self.frame_opciones, height=6, font=("Segoe UI", 10),
                                      relief="flat", borderwidth=1, padx=10, pady=10)
        self.texto_opciones.pack(fill="x")

        ttk.Label(panel, text="Descripción:", font=("Segoe UI", 10, "bold")).pack(anchor="w", pady=(10, 0))
        self.entry_descripcion = ttk.Entry(panel, width=50)
        self.entry_descripcion.pack(fill="x", pady=10)

        self.requerido_var = tk.BooleanVar(value=True)
        tk.Checkbutton(panel, text="Campo requerido", variable=self.requerido_var,
                       font=("Segoe UI", 10)).pack(anchor="w", pady=5)

        btns = ttk.Frame(cont)
        btns.pack(fill="x")

        ttk.Button(btns, text="Guardar", style="Notion.TButton",
                   command=self.guardar_campo).pack(side="left", padx=8)
        ttk.Button(btns, text="Cancelar", style="Notion.TButton",
                   command=self.ventana.destroy).pack(side="left", padx=8)

    def cargar_datos_existentes(self, campo):
        self.entry_id.insert(0, campo.get('id', ''))
        self.entry_id.config(state='disabled')
        
        self.entry_nombre.insert(0, campo.get('nombre', ''))
        self.tipo_var.set(campo.get('tipo', 'texto'))
        self.entry_descripcion.insert(0, campo.get('descripcion', ''))
        self.requerido_var.set(campo.get('requerido', False))
        
        if campo.get('tipo') == 'seleccion' and 'opciones' in campo:
            self.texto_opciones.insert("1.0", "\n".join(campo['opciones']))
    
    def guardar_campo(self):
        campo_id = self.entry_id.get().strip()
        nombre = self.entry_nombre.get().strip()
        tipo = self.tipo_var.get()
        descripcion = self.entry_descripcion.get().strip()
        requerido = self.requerido_var.get()
        
        if not campo_id:
            messagebox.showwarning("Advertencia", "El ID del campo es requerido.")
            return
        
        if not nombre:
            messagebox.showwarning("Advertencia", "El nombre del campo es requerido.")
            return
        
        if ' ' in campo_id:
            messagebox.showwarning("Advertencia", "El ID no puede contener espacios.")
            return
        
        campo = {
            'id': campo_id,
            'nombre': nombre,
            'tipo': tipo,
            'descripcion': descripcion,
            'requerido': requerido
        }
        
        if tipo == 'seleccion':
            opciones_texto = self.texto_opciones.get("1.0", tk.END).strip()
            if opciones_texto:
                campo['opciones'] = [opcion.strip() for opcion in opciones_texto.split('\n') if opcion.strip()]
            else:
                messagebox.showwarning("Advertencia", "Debe proporcionar opciones para el campo de selección.")
                return
        
        self.campo_creado = campo
        self.ventana.destroy()


def verificar_dependencias():
    try:
        from docx import Document
        return True
    except ImportError as e:
        print(f"""
        ❌ DEPENDENCIAS REQUERIDAS NO INSTALADAS
        
        Ejecute en la terminal:
        pip install python-docx
        
        Error: {e}
        """)
        return False

if __name__ == "__main__":
    app = SistemaPlantillasPersonalizadas()
    app.root.mainloop()